"""
Fail-closed reviewer for the final outputs of the "إنشاء تكست لونج" recipe.

The module deliberately separates:
1) deterministic structure and DOCX checks;
2) semantic review prompts;
3) strict validation of semantic-review responses.

No generated text is trusted to define its own schema.  The canonical order,
labels, identifiers, evidence, and response identity are all checked locally.
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import unicodedata
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from engine import EngineError, log


# The versioned JSON document is the only contract source. The reviewer,
# runtime, and copied review_contract.json all use the same fingerprint.
from long_text_contract import (
    BANNED_PHRASES,
    CANONICAL_LABELS,
    CONSUMER_EXACT,
    CONTEXT_FIELDS,
    CONTRACT,
    CONTRACT_FINGERPRINT,
    COUNTS,
    EXPECTED_TOKEN_KEYS,
    FIELD_ALIASES,
    FIELD_KEYS,
    FORMATTING,
    HEADER_ALIASES,
    REQUIRED_ARABIC_HASHTAGS,
    SEMANTIC_CROSS_CHECKS,
    SEMANTIC_ERROR_CODES,
    SEMANTIC_FIELD_KEYS,
    TOKEN_KINDS,
    TOKEN_SPECS,
)


LONG_TEXT_REVIEW_SCHEMA_VERSION = CONTRACT["semantic"]["response_schema_version"]

ARABIC_DIACRITICS_RE = re.compile(r"[\u0610-\u061a\u064b-\u065f\u0670\u06d6-\u06ed]")
INVISIBLE_CHARS_RE = re.compile(
    r"[\u200b\u200c\u200d\u200e\u200f\u202a-\u202e\u2060\u2066-\u2069\ufeff]"
)
HASHTAG_RE = re.compile(r"(?<![\w#])#[\w\u0600-\u06ff]+", re.UNICODE)
WORD_RE = re.compile(r"[0-9A-Za-zÀ-ÖØ-öø-ÿ\u0600-\u06ff]+", re.UNICODE)


# Readable presentation differences are recorded but do not make parsed data
# unusable. Structural identifiers, Heading 2 boundaries, missing/moved
# fields, counts, and consumer-exact labels remain blocking errors.
WARNING_CODES = {
    "LABEL_CHANGED",
    "DOCX_PARAGRAPH_NOT_RIGHT_ALIGNED",
    "DOCX_PARAGRAPH_RTL_MISSING",
    "DOCX_RUN_RTL_MISSING",
    "DOCX_RUN_ARABIC_LANGUAGE_MISSING",
    "DOCX_RUN_FONT_WRONG",
    "DOCX_RUN_SIZE_WRONG",
    "DOCX_LINE_SPACING_WRONG",
}

# Only labels used by a literal downstream extractor are blocking when an
# otherwise recognized alias is used.
CONSUMER_EXACT_LABEL_KEYS = frozenset(
    CONSUMER_EXACT["consumer_exact_label_keys"]
)


def _bool(value, default=False):
    if value is None:
        return default
    if isinstance(value, bool):
        return value
    return str(value).strip().lower() not in {"false", "0", "no", "off", ""}


def _clean_ws(value):
    return re.sub(r"[ \t\r\n\u00a0]+", " ", str(value or "")).strip()


def _arabic_norm(value):
    text = unicodedata.normalize("NFKC", str(value or ""))
    text = ARABIC_DIACRITICS_RE.sub("", text).replace("\u0640", "")
    text = text.translate(str.maketrans({"أ": "ا", "إ": "ا", "آ": "ا", "ى": "ي"}))
    text = re.sub(r"\s+", " ", text)
    return text.strip().casefold()


def _label_norm(value):
    text = _arabic_norm(value)
    text = text.replace("：", ":")
    text = re.sub(r"\s*:\s*", ":", text)
    text = re.sub(r"\s*[-–—]\s*", "-", text)
    text = re.sub(r"\s*\.\s*", ".", text)
    return text.strip()


def _visible_words(value):
    return WORD_RE.findall(str(value or ""))


def _paragraph_text(paragraph):
    """Read all visible text nodes, including text inside hyperlinks."""
    pieces = []
    for node in paragraph._p.iter():
        if node.tag == qn("w:t") and node.text:
            pieces.append(node.text)
        elif node.tag == qn("w:tab"):
            pieces.append("\t")
        elif node.tag in {qn("w:br"), qn("w:cr")}:
            pieces.append("\n")
    return "".join(pieces)


def _strict_json_load(path):
    def no_duplicates(pairs):
        result = {}
        for key, value in pairs:
            if key in result:
                raise ValueError(f"مفتاح JSON مكرر: {key}")
            result[key] = value
        return result

    with open(path, "r", encoding="utf-8-sig") as handle:
        return json.load(handle, object_pairs_hook=no_duplicates)


def _canonical_positive_id(value, name="id"):
    if type(value) is not int or value < 1:
        raise ValueError(f"{name} لازم يكون JSON integer موجب")
    return str(value)


def _issue(code, message, *, topic_id=None, file=None, location=None,
           severity=None, expected=None, actual=None):
    if severity is None:
        severity = "warning" if code in WARNING_CODES else "error"
    item = {
        "code": code,
        "severity": severity,
        "message": message,
    }
    if topic_id is not None:
        item["topic_id"] = int(topic_id)
    if file:
        item["file"] = file
    if location:
        item["location"] = location
    if expected is not None:
        item["expected"] = expected
    if actual is not None:
        item["actual"] = actual
    return item


def _add_unique(issues, item, seen):
    signature = json.dumps(item, ensure_ascii=False, sort_keys=True, default=str)
    if signature not in seen:
        seen.add(signature)
        issues.append(item)


def _read_topics(path, issues, seen):
    try:
        data = _strict_json_load(path)
    except (OSError, UnicodeError, json.JSONDecodeError, ValueError) as exc:
        _add_unique(
            issues,
            _issue(
                "TOPICS_INVALID_JSON",
                f"تعذر قراءة topics.json كـ JSON صارم: {exc}",
                file=os.path.basename(path),
            ),
            seen,
        )
        return {}, [], {}

    metadata = {}
    if isinstance(data, dict):
        has_titles = "titles" in data
        has_topics = "topics" in data
        if has_titles and has_topics:
            _add_unique(
                issues,
                _issue(
                    "TOPICS_AMBIGUOUS_ROOT",
                    "topics.json يحتوي titles وtopics معاً؛ ده ممكن يسقط قائمة بصمت",
                    file=os.path.basename(path),
                ),
                seen,
            )
        items = data.get("titles") if has_titles else data.get("topics")
        metadata = {key: value for key, value in data.items() if key not in {"titles", "topics"}}
    else:
        items = data

    if not isinstance(items, list):
        _add_unique(
            issues,
            _issue(
                "TOPICS_ROOT_NOT_LIST",
                "topics.json لازم يحتوي قائمة مباشرة أو مفتاح titles بقائمة",
                file=os.path.basename(path),
            ),
            seen,
        )
        return {}, [], metadata

    declared_total = metadata.get("total_count")
    if declared_total is not None and (
        type(declared_total) is not int or declared_total != len(items)
    ):
        _add_unique(
            issues,
            _issue(
                "TOPICS_TOTAL_COUNT_MISMATCH",
                "total_count لا يساوي عدد عناصر قائمة العناوين",
                file=os.path.basename(path),
                expected=len(items),
                actual=declared_total,
            ),
            seen,
        )

    titles = {}
    order = []
    normalized_title_ids = defaultdict(list)
    for index, item in enumerate(items, start=1):
        location = f"عنصر {index}"
        if not isinstance(item, dict):
            _add_unique(
                issues,
                _issue(
                    "TOPIC_ITEM_NOT_OBJECT",
                    "عنصر العنوان مش كائن JSON",
                    file=os.path.basename(path),
                    location=location,
                    actual=type(item).__name__,
                ),
                seen,
            )
            continue
        try:
            topic_id = _canonical_positive_id(item.get("id"), f"id في العنصر {index}")
        except ValueError as exc:
            _add_unique(
                issues,
                _issue(
                    "TOPIC_ID_INVALID",
                    str(exc),
                    file=os.path.basename(path),
                    location=location,
                    actual=item.get("id"),
                ),
                seen,
            )
            continue
        if topic_id in titles:
            _add_unique(
                issues,
                _issue(
                    "TOPIC_ID_DUPLICATE",
                    f"معرف الموضوع {topic_id} مكرر في topics.json",
                    topic_id=topic_id,
                    file=os.path.basename(path),
                    location=location,
                ),
                seen,
            )
            continue
        title = item.get("title")
        if not isinstance(title, str) or not title.strip():
            _add_unique(
                issues,
                _issue(
                    "TOPIC_TITLE_EMPTY",
                    f"عنوان الموضوع {topic_id} فارغ أو مش نص",
                    topic_id=topic_id,
                    file=os.path.basename(path),
                    location=location,
                ),
                seen,
            )
            title = ""
        if re.search(r"<<<\s*(?:SCRIPT|END_SCRIPT)_?\d*\s*>>>", title, re.I):
            _add_unique(
                issues,
                _issue(
                    "TOPIC_TITLE_MARKER_INJECTION",
                    "العنوان يحتوي ماركر يقدر يكسر تقسيم المخرجات",
                    topic_id=topic_id,
                    file=os.path.basename(path),
                    location=location,
                    actual=title,
                ),
                seen,
            )
        titles[topic_id] = title.strip()
        order.append(topic_id)
        if title.strip():
            normalized_title_ids[_arabic_norm(title)].append(topic_id)

    for norm_title, ids in normalized_title_ids.items():
        if len(ids) > 1:
            for topic_id in ids:
                _add_unique(
                    issues,
                    _issue(
                        "TOPIC_TITLE_DUPLICATE_NORMALIZED",
                        f"العنوان مكرر بعد التطبيع بين الموضوعات: {', '.join(ids)}",
                        topic_id=topic_id,
                        file=os.path.basename(path),
                        actual=titles.get(topic_id, ""),
                    ),
                    seen,
                )
    return titles, order, metadata


def _package_text(xml_bytes):
    try:
        text = xml_bytes.decode("utf-8", errors="replace")
    except Exception:
        return ""
    chunks = re.findall(r"<w:t(?:\s[^>]*)?>(.*?)</w:t>", text, flags=re.DOTALL)
    return _clean_ws(" ".join(re.sub(r"<[^>]+>", "", chunk) for chunk in chunks))


def _audit_docx_package(path, role, issues, seen):
    filename = os.path.basename(path)
    try:
        if not zipfile.is_zipfile(path):
            raise zipfile.BadZipFile("الملف ليس حزمة ZIP/DOCX")
        with zipfile.ZipFile(path, "r") as archive:
            names = archive.namelist()
            duplicate_names = [name for name, count in Counter(names).items() if count > 1]
            if duplicate_names:
                _add_unique(
                    issues,
                    _issue(
                        "DOCX_DUPLICATE_ZIP_ENTRIES",
                        "حزمة Word تحتوي أسماء أجزاء مكررة",
                        file=filename,
                        actual=duplicate_names,
                    ),
                    seen,
                )
            unsafe = [name for name in names if name.startswith(("/", "\\")) or ".." in Path(name).parts]
            if unsafe:
                _add_unique(
                    issues,
                    _issue(
                        "DOCX_UNSAFE_PACKAGE_PATH",
                        "حزمة Word تحتوي مسارات غير آمنة",
                        file=filename,
                        actual=unsafe,
                    ),
                    seen,
                )
            bad_member = archive.testzip()
            if bad_member:
                _add_unique(
                    issues,
                    _issue(
                        "DOCX_CORRUPT_MEMBER",
                        f"جزء تالف داخل ملف Word: {bad_member}",
                        file=filename,
                    ),
                    seen,
                )
            if "word/document.xml" not in names:
                _add_unique(
                    issues,
                    _issue(
                        "DOCX_DOCUMENT_XML_MISSING",
                        "word/document.xml غير موجود داخل ملف Word",
                        file=filename,
                    ),
                    seen,
                )
                return

            total_uncompressed = sum(info.file_size for info in archive.infolist())
            if total_uncompressed > 100 * 1024 * 1024:
                _add_unique(
                    issues,
                    _issue(
                        "DOCX_PACKAGE_TOO_LARGE",
                        "الحجم غير المضغوط لحزمة Word أكبر من 100 ميجابايت",
                        file=filename,
                        actual=total_uncompressed,
                    ),
                    seen,
                )

            macro_parts = [name for name in names if "vbaProject" in name or name.endswith(".bin")]
            if macro_parts:
                _add_unique(
                    issues,
                    _issue(
                        "DOCX_MACRO_OR_BINARY_PART",
                        "ملف المخرجات يحتوي ماكرو أو جزءا ثنائيا غير متوقع",
                        file=filename,
                        actual=macro_parts,
                    ),
                    seen,
                )

            document_xml = archive.read("word/document.xml")
            xml_text = document_xml.decode("utf-8", errors="replace")
            structural_tags = {
                "DOCX_TRACKED_INSERTIONS": r"<w:ins(?:\s|>)",
                "DOCX_TRACKED_DELETIONS": r"<w:del(?:\s|>)",
                "DOCX_COMMENTS_ANCHORS": r"<w:commentRangeStart(?:\s|>)",
                "DOCX_TEXTBOX_CONTENT": r"<w:txbxContent(?:\s|>)",
                "DOCX_CONTENT_CONTROLS": r"<w:sdt(?:\s|>)",
                "DOCX_ALT_CHUNK": r"<w:altChunk(?:\s|>)",
                "DOCX_FIELDS": r"<w:(?:fldSimple|instrText)(?:\s|>)",
                "DOCX_DRAWING_OR_OBJECT": r"<w:(?:drawing|object|pict)(?:\s|>)",
            }
            for code, pattern in structural_tags.items():
                if re.search(pattern, xml_text):
                    _add_unique(
                        issues,
                        _issue(
                            code,
                            "ملف Word يحتوي بنية غير متوقعة قد لا يقرأها التجميع بنفس الشكل",
                            file=filename,
                            location="word/document.xml",
                        ),
                        seen,
                    )

            non_body_parts = [
                name for name in names
                if re.fullmatch(r"word/(?:header|footer)\d+\.xml", name)
                or name in {"word/footnotes.xml", "word/endnotes.xml", "word/comments.xml"}
            ]
            for part_name in non_body_parts:
                part_text = _package_text(archive.read(part_name))
                if part_text:
                    _add_unique(
                        issues,
                        _issue(
                            "DOCX_TEXT_OUTSIDE_BODY",
                            "يوجد نص خارج جسم المستند وقد تتجاهله الوصفة التالية",
                            file=filename,
                            location=part_name,
                            actual=part_text[:300],
                        ),
                        seen,
                    )

            if "word/settings.xml" in names:
                settings = archive.read("word/settings.xml").decode("utf-8", errors="replace")
                if re.search(r"<w:documentProtection(?:\s|>)", settings):
                    _add_unique(
                        issues,
                        _issue(
                            "DOCX_PROTECTED",
                            "ملف Word محمي ضد التعديل أو القراءة الكاملة",
                            file=filename,
                            location="word/settings.xml",
                        ),
                        seen,
                    )
    except (OSError, zipfile.BadZipFile, RuntimeError) as exc:
        _add_unique(
            issues,
            _issue(
                "DOCX_INVALID_PACKAGE",
                f"تعذر فتح ملف Word كحزمة سليمة: {exc}",
                file=filename,
            ),
            seen,
        )


def _run_has_text(run_element):
    return bool(_clean_ws("".join(node.text or "" for node in run_element.iter(qn("w:t")))))


def _paragraph_format_issues(
    paragraph,
    *,
    paragraph_index,
    role,
    is_script_heading,
    topic_id,
    expected_font,
    expected_size,
    expected_line_spacing,
    issues,
    seen,
):
    filename = role
    location = f"فقرة {paragraph_index + 1}"
    expected_style = "Heading 2" if is_script_heading else "Normal"
    style_name = paragraph.style.name if paragraph.style is not None else ""
    if style_name != expected_style:
        _add_unique(
            issues,
            _issue(
                "DOCX_PARAGRAPH_STYLE_WRONG",
                f"نمط الفقرة لازم يكون {expected_style}",
                topic_id=topic_id,
                file=filename,
                location=location,
                expected=expected_style,
                actual=style_name,
            ),
            seen,
        )

    if paragraph.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
        _add_unique(
            issues,
            _issue(
                "DOCX_PARAGRAPH_NOT_RIGHT_ALIGNED",
                "الفقرة ليست بمحاذاة يمين صريحة",
                topic_id=topic_id,
                file=filename,
                location=location,
                expected="RIGHT",
                actual=str(paragraph.alignment),
            ),
            seen,
        )

    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    bidi_value = bidi.get(qn("w:val")) if bidi is not None else None
    # The production save_docx path deliberately writes w:bidi=0 together
    # with explicit right alignment and run-level RTL because that is the
    # representation Word COM reports correctly in this application. Missing
    # w:bidi is still reported; right alignment and every run are checked
    # independently below.
    if bidi is None:
        _add_unique(
            issues,
            _issue(
                "DOCX_PARAGRAPH_RTL_MISSING",
                "خاصية BIDI/RTL غير موجودة على مستوى الفقرة",
                topic_id=topic_id,
                file=filename,
                location=location,
                expected="w:bidi موجود مع محاذاة يمين وRTL على كل Run",
                actual="missing",
            ),
            seen,
        )

    if not is_script_heading and expected_line_spacing:
        spacing = p_pr.find(qn("w:spacing"))
        expected_twips = str(int(round(float(expected_line_spacing) * 20)))
        actual_line = spacing.get(qn("w:line")) if spacing is not None else None
        if actual_line != expected_twips:
            _add_unique(
                issues,
                _issue(
                    "DOCX_LINE_SPACING_WRONG",
                    "المسافة بين السطور لا تطابق عقد المخرج",
                    topic_id=topic_id,
                    file=filename,
                    location=location,
                    expected=f"{expected_line_spacing} pt",
                    actual=actual_line,
                ),
                seen,
            )

    for run_index, run_element in enumerate(paragraph._p.iter(qn("w:r")), start=1):
        if not _run_has_text(run_element):
            continue
        run_location = f"{location} / run {run_index}"
        r_pr = run_element.find(qn("w:rPr"))
        rtl = r_pr.find(qn("w:rtl")) if r_pr is not None else None
        rtl_value = rtl.get(qn("w:val")) if rtl is not None else None
        if rtl is None or str(rtl_value or "1").lower() in {"0", "false", "off"}:
            _add_unique(
                issues,
                _issue(
                    "DOCX_RUN_RTL_MISSING",
                    "خاصية RTL غير مفعلة على مستوى النص Run",
                    topic_id=topic_id,
                    file=filename,
                    location=run_location,
                ),
                seen,
            )
        lang = r_pr.find(qn("w:lang")) if r_pr is not None else None
        bidi_lang = lang.get(qn("w:bidi")) if lang is not None else None
        if not bidi_lang or not str(bidi_lang).lower().startswith("ar"):
            _add_unique(
                issues,
                _issue(
                    "DOCX_RUN_ARABIC_LANGUAGE_MISSING",
                    "لغة النص العربي غير مضبوطة على مستوى Run",
                    topic_id=topic_id,
                    file=filename,
                    location=run_location,
                    expected="ar-*",
                    actual=bidi_lang,
                ),
                seen,
            )
        if is_script_heading:
            continue
        fonts = r_pr.find(qn("w:rFonts")) if r_pr is not None else None
        font_values = {
            attr: fonts.get(qn(f"w:{attr}")) if fonts is not None else None
            for attr in ("ascii", "hAnsi", "cs")
        }
        if any(value != expected_font for value in font_values.values()):
            _add_unique(
                issues,
                _issue(
                    "DOCX_RUN_FONT_WRONG",
                    "خط الـ Run لا يطابق الخط المطلوب",
                    topic_id=topic_id,
                    file=filename,
                    location=run_location,
                    expected=expected_font,
                    actual=font_values,
                ),
                seen,
            )
        sz_cs = r_pr.find(qn("w:szCs")) if r_pr is not None else None
        actual_size = sz_cs.get(qn("w:val")) if sz_cs is not None else None
        expected_half_points = str(int(round(float(expected_size) * 2)))
        if actual_size != expected_half_points:
            _add_unique(
                issues,
                _issue(
                    "DOCX_RUN_SIZE_WRONG",
                    "حجم خط الـ Run لا يطابق الحجم المطلوب",
                    topic_id=topic_id,
                    file=filename,
                    location=run_location,
                    expected=f"{expected_size} pt",
                    actual=actual_size,
                ),
                seen,
            )


def _parse_script_header(text):
    raw = _clean_ws(text)
    canonical = re.fullmatch(r"Script ([1-9]\d*)", raw)
    if canonical:
        return canonical.group(1), True
    translated_digits = raw.translate(str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789"))
    alias = re.fullmatch(
        r"(?:script|سكريبت)\s*[:#_\-]?\s*(0*\d+)",
        translated_digits,
        flags=re.IGNORECASE,
    )
    if alias and int(alias.group(1)) > 0:
        return str(int(alias.group(1))), False
    return None, False


def _load_document_topics(
    path,
    *,
    role,
    expected_font,
    expected_size,
    expected_line_spacing,
    issues,
    seen,
):
    filename = os.path.basename(path)
    _audit_docx_package(path, role, issues, seen)
    try:
        document = Document(path)
    except Exception as exc:
        _add_unique(
            issues,
            _issue(
                "DOCX_OPEN_FAILED",
                f"تعذر فتح ملف Word: {exc}",
                file=filename,
            ),
            seen,
        )
        return [], document if "document" in locals() else None

    for table_index, table in enumerate(document.tables, start=1):
        table_text = _clean_ws(
            " ".join(cell.text for row in table.rows for cell in row.cells)
        )
        if table_text:
            _add_unique(
                issues,
                _issue(
                    "DOCX_TABLE_CONTENT",
                    "محتوى داخل جدول Word لن تقرأه الوصفة التالية كفقرات عادية",
                    file=filename,
                    location=f"جدول {table_index}",
                    actual=table_text[:500],
                ),
                seen,
            )

    entries = []
    current = None
    occurrence_count = Counter()
    for paragraph_index, paragraph in enumerate(document.paragraphs):
        raw_text = _paragraph_text(paragraph)
        stripped = raw_text.strip()
        parsed_id, canonical_header = _parse_script_header(stripped) if stripped else (None, False)
        if parsed_id is not None:
            occurrence_count[parsed_id] += 1
            current = {
                "id": parsed_id,
                "heading_index": paragraph_index,
                "heading_text": stripped,
                "paragraphs": [],
                "occurrence": occurrence_count[parsed_id],
            }
            entries.append(current)
            if not canonical_header:
                _add_unique(
                    issues,
                    _issue(
                        "SCRIPT_HEADER_CHANGED",
                        "عنوان الموضوع في Word ليس بالصيغة الحرفية Script N",
                        topic_id=parsed_id,
                        file=filename,
                        location=f"فقرة {paragraph_index + 1}",
                        expected=f"Script {parsed_id}",
                        actual=stripped,
                    ),
                    seen,
                )
            if occurrence_count[parsed_id] > 1:
                _add_unique(
                    issues,
                    _issue(
                        "SCRIPT_ID_DUPLICATE",
                        f"Script {parsed_id} مكرر داخل الملف",
                        topic_id=parsed_id,
                        file=filename,
                        location=f"فقرة {paragraph_index + 1}",
                    ),
                    seen,
                )
            _paragraph_format_issues(
                paragraph,
                paragraph_index=paragraph_index,
                role=filename,
                is_script_heading=True,
                topic_id=parsed_id,
                expected_font=expected_font,
                expected_size=expected_size,
                expected_line_spacing=expected_line_spacing,
                issues=issues,
                seen=seen,
            )
            continue

        if stripped and re.match(r"^(?:Script|سكريبت)\b", stripped, flags=re.IGNORECASE):
            _add_unique(
                issues,
                _issue(
                    "SCRIPT_HEADER_MALFORMED",
                    "فقرة تشبه عنوان Script لكن رقمها أو صيغتها غير قابلة للقراءة الآمنة",
                    file=filename,
                    location=f"فقرة {paragraph_index + 1}",
                    actual=stripped,
                ),
                seen,
            )
            current = None
            continue

        if current is None:
            if stripped:
                _add_unique(
                    issues,
                    _issue(
                        "DOCX_ORPHAN_CONTENT",
                        "محتوى موجود خارج أي Script",
                        file=filename,
                        location=f"فقرة {paragraph_index + 1}",
                        actual=stripped[:500],
                    ),
                    seen,
                )
            continue

        if not stripped:
            _add_unique(
                issues,
                _issue(
                    "DOCX_EMPTY_PARAGRAPH",
                    "فقرة فارغة داخل الموضوع تكسر العقد الحرفي للمخرج",
                    topic_id=current["id"],
                    file=filename,
                    location=f"فقرة {paragraph_index + 1}",
                ),
                seen,
            )
            continue

        current["paragraphs"].append(
            {
                "index": paragraph_index,
                "text": raw_text,
                "paragraph": paragraph,
            }
        )
        _paragraph_format_issues(
            paragraph,
            paragraph_index=paragraph_index,
            role=filename,
            is_script_heading=False,
            topic_id=current["id"],
            expected_font=expected_font,
            expected_size=expected_size,
            expected_line_spacing=expected_line_spacing,
            issues=issues,
            seen=seen,
        )

    if not entries:
        _add_unique(
            issues,
            _issue(
                "DOCX_NO_SCRIPTS",
                "ملف Word لا يحتوي أي عنوان Script صالح",
                file=filename,
            ),
            seen,
        )
    return entries, document


def _header_lookup():
    lookup = {}
    for key, aliases in HEADER_ALIASES.items():
        for alias in aliases:
            lookup[_label_norm(alias)] = key
    return lookup


HEADER_LOOKUP = _header_lookup()


def _match_header(line):
    normalized = _label_norm(line)
    key = HEADER_LOOKUP.get(normalized)
    if key:
        return key, line.strip()
    return None, None


def _split_label_value(line):
    match = re.match(r"^\s*(.{1,80}?)([:：])\s*(.*)$", line)
    if match:
        return match.group(1).strip(), match.group(3)
    return line.strip(), ""


def _field_match_score(raw_label, alias):
    left = _label_norm(raw_label)
    right = _label_norm(alias)
    if left == right:
        return 100
    if left.replace(" ", "") == right.replace(" ", ""):
        return 96
    # Small deterministic typo tolerance without accepting unrelated prose.
    from difflib import SequenceMatcher

    return int(SequenceMatcher(None, left, right).ratio() * 100)


def _match_field(line, context):
    raw_label, value = _split_label_value(line)
    candidate_keys = CONTEXT_FIELDS.get(context, ())
    best = None
    for key in candidate_keys:
        for alias in FIELD_ALIASES[key]:
            score = _field_match_score(raw_label, alias)
            if best is None or score > best[0]:
                best = (score, key, alias)
    if best and best[0] >= 84:
        return best[1], raw_label, value

    # Unique long labels can still be assigned when they appear in the wrong
    # section; generic "العنوان/الوصف" cannot.
    for key in (
        "youtube_title_1",
        "youtube_title_2",
        "youtube_description_1",
        "youtube_description_2",
        "youtube_thumbnail_1",
        "youtube_thumbnail_2",
        "tiktok_screen",
        "facebook_thumbnail",
    ):
        for alias in FIELD_ALIASES[key]:
            if _field_match_score(raw_label, alias) >= 90:
                return key, raw_label, value
    return None, raw_label, value


def _looks_like_unknown_label(line):
    prefix, _value = _split_label_value(line)
    norm = _arabic_norm(prefix)
    return bool(
        re.match(
            r"^(?:القسم|العنوان|عنوان|الوصف|وصف|جمله|الكلمات|كلمات|الترجمه|title|description)\b",
            norm,
            flags=re.IGNORECASE,
        )
    )


def _parse_script_fields(entry, filename, issues, seen):
    topic_id = entry["id"]
    sequence = []
    occurrences = defaultdict(list)
    context = None
    current_field = None
    current_lines = []
    current_location = None

    def finalize():
        nonlocal current_field, current_lines, current_location
        if current_field is None:
            return
        value = "\n".join(current_lines).strip()
        occurrences[current_field].append(
            {"value": value, "location": current_location}
        )
        if not value:
            _add_unique(
                issues,
                _issue(
                    "FIELD_VALUE_EMPTY",
                    f"قيمة البند {CANONICAL_LABELS[current_field]} فارغة",
                    topic_id=topic_id,
                    file=filename,
                    location=current_location,
                ),
                seen,
            )
        current_field = None
        current_lines = []
        current_location = None

    for paragraph_record in entry["paragraphs"]:
        paragraph_index = paragraph_record["index"]
        text = paragraph_record["text"]
        for line_index, raw_line in enumerate(text.splitlines() or [text], start=1):
            line = raw_line.strip()
            location = f"فقرة {paragraph_index + 1} / سطر {line_index}"
            if not line:
                if current_field is not None:
                    current_lines.append("")
                continue

            header_key, raw_header = _match_header(line)
            if header_key:
                finalize()
                sequence.append(header_key)
                context = header_key
                if header_key == "section_translations":
                    context = "section_translations"
                canonical = CANONICAL_LABELS[header_key]
                if raw_header != canonical:
                    _add_unique(
                        issues,
                        _issue(
                            "LABEL_CHANGED",
                            "صيغة عنوان القسم أو الترجمة اتغيرت عن العقد الحرفي",
                            topic_id=topic_id,
                            file=filename,
                            location=location,
                            expected=canonical,
                            actual=raw_header,
                            severity="warning",
                        ),
                        seen,
                    )
                continue

            field_key, raw_label, inline_value = _match_field(line, context)
            if field_key:
                finalize()
                sequence.append(field_key)
                current_field = field_key
                current_location = location
                current_lines = [inline_value] if inline_value else []
                canonical_label = CANONICAL_LABELS[field_key].rstrip(":")
                actual_with_colon = raw_label + (":" if ":" in line or "：" in line else "")
                if actual_with_colon != CANONICAL_LABELS[field_key]:
                    _add_unique(
                        issues,
                        _issue(
                            "LABEL_CHANGED",
                            "اسم البند أو علامته اتغيرت عن الصيغة الحرفية المطلوبة",
                            topic_id=topic_id,
                            file=filename,
                            location=location,
                            expected=CANONICAL_LABELS[field_key],
                            actual=actual_with_colon or canonical_label,
                            severity=(
                                "error"
                                if field_key in CONSUMER_EXACT_LABEL_KEYS
                                else "warning"
                            ),
                        ),
                        seen,
                    )
                if field_key not in CONTEXT_FIELDS.get(context, ()):
                    _add_unique(
                        issues,
                        _issue(
                            "FIELD_IN_WRONG_SECTION",
                            "البند موجود تحت قسم غير القسم المخصص له",
                            topic_id=topic_id,
                            file=filename,
                            location=location,
                            expected=field_key,
                            actual=context,
                        ),
                        seen,
                    )
                continue

            if _looks_like_unknown_label(line):
                _add_unique(
                    issues,
                    _issue(
                        "UNKNOWN_OR_MALFORMED_LABEL",
                        "سطر يشبه اسما لقسم أو بند لكنه غير قابل للمطابقة الآمنة",
                        topic_id=topic_id,
                        file=filename,
                        location=location,
                        actual=line[:300],
                    ),
                    seen,
                )
            if current_field is None:
                _add_unique(
                    issues,
                    _issue(
                        "CONTENT_OUTSIDE_FIELD",
                        "محتوى موجود داخل Script لكن خارج أي بند معروف",
                        topic_id=topic_id,
                        file=filename,
                        location=location,
                        actual=line[:500],
                    ),
                    seen,
                )
            else:
                current_lines.append(raw_line.rstrip())
    finalize()

    counts = Counter(sequence)
    for key in EXPECTED_TOKEN_KEYS:
        count = counts.get(key, 0)
        if count == 0:
            _add_unique(
                issues,
                _issue(
                    "TOKEN_MISSING",
                    f"القسم أو البند مفقود: {CANONICAL_LABELS[key]}",
                    topic_id=topic_id,
                    file=filename,
                    expected=key,
                ),
                seen,
            )
        elif count > 1:
            _add_unique(
                issues,
                _issue(
                    "TOKEN_DUPLICATE",
                    f"القسم أو البند مكرر {count} مرات: {CANONICAL_LABELS[key]}",
                    topic_id=topic_id,
                    file=filename,
                    expected=key,
                    actual=count,
                ),
                seen,
            )

    unknown_tokens = [key for key in sequence if key not in EXPECTED_TOKEN_KEYS]
    if unknown_tokens:
        _add_unique(
            issues,
            _issue(
                "TOKEN_UNKNOWN",
                "تم التعرف على توكنات غير موجودة في العقد",
                topic_id=topic_id,
                file=filename,
                actual=unknown_tokens,
            ),
            seen,
        )

    if sequence != list(EXPECTED_TOKEN_KEYS):
        _add_unique(
            issues,
            _issue(
                "TOKEN_SEQUENCE_WRONG",
                "ترتيب الأقسام أو البنود لا يطابق الترتيب الإلزامي",
                topic_id=topic_id,
                file=filename,
                expected=list(EXPECTED_TOKEN_KEYS),
                actual=sequence,
            ),
            seen,
        )
        previous_index = -1
        for key in sequence:
            if key not in EXPECTED_TOKEN_KEYS:
                continue
            expected_index = EXPECTED_TOKEN_KEYS.index(key)
            if expected_index < previous_index:
                _add_unique(
                    issues,
                    _issue(
                        "TOKEN_MOVED_EARLIER",
                        f"البند {CANONICAL_LABELS[key]} متقدم عن مكانه الصحيح",
                        topic_id=topic_id,
                        file=filename,
                        actual=key,
                    ),
                    seen,
                )
            previous_index = max(previous_index, expected_index)

    fields = {}
    for key in FIELD_KEYS:
        values = occurrences.get(key, [])
        if len(values) == 1:
            fields[key] = values[0]["value"]
        elif values:
            fields[key] = values[0]["value"]
    return fields, sequence


def _extract_hashtags(value):
    return HASHTAG_RE.findall(str(value or ""))


def _strip_hashtags(value):
    return _clean_ws(HASHTAG_RE.sub("", str(value or "")))


def _validate_thumbnail_value(
    value,
    *,
    key,
    topic_id,
    filename,
    min_words,
    max_words,
    issues,
    seen,
):
    lines = [line.strip() for line in str(value or "").splitlines() if line.strip()]
    words = _visible_words(value)
    if not min_words <= len(words) <= max_words:
        _add_unique(
            issues,
            _issue(
                "SHORT_PHRASE_WORD_COUNT",
                f"{CANONICAL_LABELS[key]} فيها {len(words)} كلمة",
                topic_id=topic_id,
                file=filename,
                expected=f"{min_words}-{max_words}",
                actual=len(words),
            ),
            seen,
        )
    if not COUNTS["short_phrase_min_lines"] <= len(lines) <= COUNTS["short_phrase_max_lines"]:
        _add_unique(
            issues,
            _issue(
                "SHORT_PHRASE_LINE_COUNT",
                f"{CANONICAL_LABELS[key]} لازم تتقسم على سطرين أو 3",
                topic_id=topic_id,
                file=filename,
                expected=f"{COUNTS['short_phrase_min_lines']}-{COUNTS['short_phrase_max_lines']}",
                actual=len(lines),
            ),
            seen,
        )
    for line_index, line in enumerate(lines, start=1):
        count = len(_visible_words(line))
        if count > COUNTS["short_phrase_max_words_per_line"]:
            _add_unique(
                issues,
                _issue(
                    "SHORT_PHRASE_TOO_MANY_WORDS_PER_LINE",
                    "سطر في الجملة القصيرة يحتوي أكثر من 3 كلمات",
                    topic_id=topic_id,
                    file=filename,
                    location=f"{CANONICAL_LABELS[key]} / سطر {line_index}",
                    expected=f"<={COUNTS['short_phrase_max_words_per_line']}",
                    actual=count,
                ),
                seen,
            )


def _validate_keyword_list(value, *, key, topic_id, filename, issues, seen):
    raw = str(value or "").strip()
    delimiter = COUNTS["keywords_delimiter"]
    if "،" in raw:
        _add_unique(
            issues,
            _issue(
                "KEYWORDS_WRONG_DELIMITER",
                "الكلمات المفتاحية لا تستخدم الفاصل المحدد في العقد",
                topic_id=topic_id,
                file=filename,
                expected=delimiter,
                actual="،",
            ),
            seen,
        )
    items = [item.strip() for item in raw.split(delimiter) if item.strip()]
    if len(items) != COUNTS["keywords_each"]:
        _add_unique(
            issues,
            _issue(
                "KEYWORDS_COUNT_WRONG",
                f"{CANONICAL_LABELS[key]} تحتوي {len(items)} عنصرا",
                topic_id=topic_id,
                file=filename,
                expected=COUNTS["keywords_each"],
                actual=len(items),
            ),
            seen,
        )
    normalized = [_arabic_norm(item) for item in items]
    duplicates = sorted(
        {items[index] for index, item in enumerate(normalized) if normalized.count(item) > 1}
    )
    if duplicates:
        _add_unique(
            issues,
            _issue(
                "KEYWORDS_DUPLICATE",
                "قائمة الكلمات المفتاحية تحتوي عناصر مكررة",
                topic_id=topic_id,
                file=filename,
                actual=duplicates,
            ),
            seen,
        )


def _validate_description_hashtags(
    value,
    *,
    key,
    topic_id,
    filename,
    require_arabic_defaults,
    issues,
    seen,
):
    hashtags = _extract_hashtags(value)
    if len(hashtags) != COUNTS["description_hashtags"]:
        _add_unique(
            issues,
            _issue(
                "HASHTAG_COUNT_WRONG",
                f"{CANONICAL_LABELS[key]} تحتوي {len(hashtags)} هاشتاج",
                topic_id=topic_id,
                file=filename,
                expected=COUNTS["description_hashtags"],
                actual=len(hashtags),
            ),
            seen,
        )
    normalized = [_arabic_norm(tag) for tag in hashtags]
    if len(set(normalized)) != len(normalized):
        _add_unique(
            issues,
            _issue(
                "HASHTAG_DUPLICATE",
                "الوصف يحتوي هاشتاجا مكررا",
                topic_id=topic_id,
                file=filename,
                actual=hashtags,
            ),
            seen,
        )
    if require_arabic_defaults:
        missing = sorted(REQUIRED_ARABIC_HASHTAGS - set(hashtags))
        if missing:
            _add_unique(
                issues,
                _issue(
                    "REQUIRED_HASHTAG_MISSING",
                    "هاشتاجات القناة الإلزامية ناقصة",
                    topic_id=topic_id,
                    file=filename,
                    expected=sorted(REQUIRED_ARABIC_HASHTAGS),
                    actual=hashtags,
                ),
                seen,
            )


def _validate_field_content(fields, title, all_titles, topic_id, filename, issues, seen):
    title_1 = fields.get("youtube_title_1", "")
    if title_1 != title:
        actual_norm = _arabic_norm(title_1)
        other_ids = [
            other_id for other_id, other_title in all_titles.items()
            if other_id != topic_id and _arabic_norm(other_title) == actual_norm
        ]
        code = "ORIGINAL_TITLE_FROM_OTHER_TOPIC" if other_ids else "ORIGINAL_TITLE_CHANGED"
        message = (
            f"العنوان الأول يخص موضوعا آخر: {', '.join(other_ids)}"
            if other_ids
            else "العنوان الأول ليس نسخة حرفية من العنوان الأصلي"
        )
        _add_unique(
            issues,
            _issue(
                code,
                message,
                topic_id=topic_id,
                file=filename,
                expected=title,
                actual=title_1,
            ),
            seen,
        )

    title_2 = fields.get("youtube_title_2", "")
    if title_2 and _arabic_norm(title_2) == _arabic_norm(title_1):
        _add_unique(
            issues,
            _issue(
                "SECOND_TITLE_NOT_DISTINCT",
                "العنوان الثاني مطابق للأول بعد التطبيع، والمطلوب كلمات مختلفة",
                topic_id=topic_id,
                file=filename,
                actual=title_2,
            ),
            seen,
        )

    description_keys = (
        "youtube_description_1",
        "youtube_description_2",
        "tiktok_description",
        "facebook_description",
        "translation_en_description",
        "translation_fr_description",
        "translation_es_description",
        "translation_de_description",
    )
    for key in description_keys:
        _validate_description_hashtags(
            fields.get(key, ""),
            key=key,
            topic_id=topic_id,
            filename=filename,
            require_arabic_defaults=key in {
                "youtube_description_1",
                "youtube_description_2",
                "tiktok_description",
                "facebook_description",
            },
            issues=issues,
            seen=seen,
        )

    first_tags = set(_extract_hashtags(fields.get("youtube_description_1", "")))
    second_tags = set(_extract_hashtags(fields.get("youtube_description_2", "")))
    if first_tags and second_tags and (
        first_tags & second_tags != REQUIRED_ARABIC_HASHTAGS
        or len(first_tags - second_tags) != COUNTS["youtube_unique_hashtags_each"]
        or len(second_tags - first_tags) != COUNTS["youtube_unique_hashtags_each"]
    ):
        _add_unique(
            issues,
            _issue(
                "YOUTUBE_DESCRIPTION_HASHTAG_DIFFERENCE_WRONG",
                "وصفي يوتيوب لازم يشتركا في 4 هاشتاجات القناة ويختلفا في الهاشتاجين الآخرين",
                topic_id=topic_id,
                file=filename,
                expected={
                    "shared": sorted(REQUIRED_ARABIC_HASHTAGS),
                    "unique_each": COUNTS["youtube_unique_hashtags_each"],
                },
                actual={
                    "first": sorted(first_tags),
                    "second": sorted(second_tags),
                },
            ),
            seen,
        )

    if (
        fields.get("youtube_description_1")
        and _arabic_norm(_strip_hashtags(fields["youtube_description_1"]))
        == _arabic_norm(_strip_hashtags(fields.get("youtube_description_2", "")))
    ):
        _add_unique(
            issues,
            _issue(
                "YOUTUBE_DESCRIPTIONS_IDENTICAL",
                "وصـفا يوتيوب متطابقان بعد حذف الهاشتاجات",
                topic_id=topic_id,
                file=filename,
            ),
            seen,
        )

    _validate_thumbnail_value(
        fields.get("youtube_thumbnail_1", ""),
        key="youtube_thumbnail_1",
        topic_id=topic_id,
        filename=filename,
        min_words=COUNTS["youtube_thumbnail_min_words"],
        max_words=COUNTS["youtube_thumbnail_max_words"],
        issues=issues,
        seen=seen,
    )
    _validate_thumbnail_value(
        fields.get("youtube_thumbnail_2", ""),
        key="youtube_thumbnail_2",
        topic_id=topic_id,
        filename=filename,
        min_words=COUNTS["youtube_thumbnail_min_words"],
        max_words=COUNTS["youtube_thumbnail_max_words"],
        issues=issues,
        seen=seen,
    )
    _validate_thumbnail_value(
        fields.get("tiktok_screen", ""),
        key="tiktok_screen",
        topic_id=topic_id,
        filename=filename,
        min_words=COUNTS["tiktok_screen_min_words"],
        max_words=COUNTS["tiktok_screen_max_words"],
        issues=issues,
        seen=seen,
    )

    for key in ("youtube_keywords", "facebook_keywords"):
        _validate_keyword_list(
            fields.get(key, ""),
            key=key,
            topic_id=topic_id,
            filename=filename,
            issues=issues,
            seen=seen,
        )

    for title_key, description_key in (
        ("tiktok_title", "tiktok_description"),
        ("facebook_title", "facebook_description"),
    ):
        platform_title = fields.get(title_key, "").strip()
        description_body = fields.get(description_key, "").lstrip()
        if platform_title and not description_body.startswith(platform_title):
            _add_unique(
                issues,
                _issue(
                    "PLATFORM_DESCRIPTION_NOT_STARTING_WITH_TITLE",
                    "وصف المنصة لا يبدأ بعنوان المنصة حرفيا",
                    topic_id=topic_id,
                    file=filename,
                    expected=platform_title,
                    actual=description_body[: len(platform_title) + 80],
                ),
                seen,
            )

    for key, value in fields.items():
        if INVISIBLE_CHARS_RE.search(value):
            chars = sorted({f"U+{ord(char):04X}" for char in value if INVISIBLE_CHARS_RE.match(char)})
            _add_unique(
                issues,
                _issue(
                    "FIELD_INVISIBLE_CHARACTERS",
                    "البند يحتوي حروف اتجاه أو حروف صفرية مخفية",
                    topic_id=topic_id,
                    file=filename,
                    location=key,
                    actual=chars,
                ),
                seen,
            )
        if "\ufffd" in value:
            _add_unique(
                issues,
                _issue(
                    "FIELD_REPLACEMENT_CHARACTER",
                    "البند يحتوي حرف الاستبدال الدال على ترميز فاسد",
                    topic_id=topic_id,
                    file=filename,
                    location=key,
                ),
                seen,
            )
        if re.search(r"<<<[^>]+>>>", value):
            _add_unique(
                issues,
                _issue(
                    "FIELD_MARKER_LEAK",
                    "ماركر MG Ranner تسرب داخل قيمة بند",
                    topic_id=topic_id,
                    file=filename,
                    location=key,
                    actual=value[:300],
                ),
                seen,
            )
        if "**" in value or "__" in value or "`" in value:
            _add_unique(
                issues,
                _issue(
                    "FIELD_MARKDOWN_LEAK",
                    "بقايا Markdown موجودة داخل البند",
                    topic_id=topic_id,
                    file=filename,
                    location=key,
                    actual=value[:300],
                ),
                seen,
            )
        normalized_value = _arabic_norm(value)
        found_banned = [
            phrase for phrase in BANNED_PHRASES
            if _arabic_norm(phrase) in normalized_value
        ]
        if found_banned:
            _add_unique(
                issues,
                _issue(
                    "BANNED_PHRASE",
                    "البند يحتوي كلمة أو تعبيرا محظورا",
                    topic_id=topic_id,
                    file=filename,
                    location=key,
                    actual=sorted(set(found_banned)),
                ),
                seen,
            )


def _validate_id_coverage(
    entries,
    titles,
    expected_ids,
    filename,
    issues,
    seen,
):
    actual_ids = [entry["id"] for entry in entries]
    actual_unique = list(dict.fromkeys(actual_ids))
    if actual_ids != actual_unique:
        # Per-ID duplicate issues were already emitted; keep one file-level fact.
        _add_unique(
            issues,
            _issue(
                "SCRIPT_ID_LIST_HAS_DUPLICATES",
                "قائمة Script في الملف تحتوي معرفات مكررة",
                file=filename,
                actual=actual_ids,
            ),
            seen,
        )
    missing = [topic_id for topic_id in expected_ids if topic_id not in actual_unique]
    extra = [topic_id for topic_id in actual_unique if topic_id not in expected_ids]
    if missing:
        _add_unique(
            issues,
            _issue(
                "SCRIPT_IDS_MISSING",
                "موضوعات مطلوبة مفقودة بالكامل من المخرج",
                file=filename,
                expected=expected_ids,
                actual={"missing": missing},
            ),
            seen,
        )
        for topic_id in missing:
            _add_unique(
                issues,
                _issue(
                    "SCRIPT_ID_MISSING",
                    f"Script {topic_id} مفقود من المخرج",
                    topic_id=topic_id,
                    file=filename,
                ),
                seen,
            )
    if extra:
        _add_unique(
            issues,
            _issue(
                "SCRIPT_IDS_EXTRA",
                "المخرج يحتوي موضوعات غير مطلوبة",
                file=filename,
                expected=expected_ids,
                actual={"extra": extra},
            ),
            seen,
        )
        for topic_id in extra:
            _add_unique(
                issues,
                _issue(
                    "SCRIPT_ID_EXTRA",
                    f"Script {topic_id} غير موجود ضمن الموضوعات المطلوبة",
                    topic_id=topic_id,
                    file=filename,
                ),
                seen,
            )
    if actual_unique != expected_ids:
        _add_unique(
            issues,
            _issue(
                "SCRIPT_ID_ORDER_WRONG",
                "ترتيب Script لا يطابق ترتيب الموضوعات المطلوبة",
                file=filename,
                expected=expected_ids,
                actual=actual_unique,
            ),
            seen,
        )
    for topic_id in expected_ids if titles is not None else ():
        if topic_id not in titles:
            _add_unique(
                issues,
                _issue(
                    "SELECTED_TOPIC_NOT_IN_TOPICS_JSON",
                    f"الموضوع المختار {topic_id} غير موجود في topics.json",
                    topic_id=topic_id,
                    file=filename,
                ),
                seen,
            )


def _validate_thumbnail_doc(
    entries,
    expected_ids,
    parsed_topics,
    filename,
    issues,
    seen,
):
    _validate_id_coverage(entries, None, expected_ids, filename, issues, seen)
    for entry in entries:
        topic_id = entry["id"]
        if len(entry["paragraphs"]) != 1:
            _add_unique(
                issues,
                _issue(
                    "THUMBNAIL_PARAGRAPH_COUNT_WRONG",
                    "كل Script في ملف الصور المصغرة لازم يتبعه فقرة واحدة بالضبط",
                    topic_id=topic_id,
                    file=filename,
                    expected=1,
                    actual=len(entry["paragraphs"]),
                ),
                seen,
            )
        actual = "\n".join(record["text"] for record in entry["paragraphs"]).strip()
        expected = (
            parsed_topics.get(topic_id, {})
            .get("fields", {})
            .get("youtube_thumbnail_1", "")
            .strip()
        )
        if actual != expected:
            _add_unique(
                issues,
                _issue(
                    "THUMBNAIL_OUTPUT_MISMATCH",
                    "جملة ملف الصور المصغرة لا تطابق جملة الفيديو الأول في scripts_output.docx",
                    topic_id=topic_id,
                    file=filename,
                    expected=expected,
                    actual=actual,
                ),
                seen,
            )


def _request_id(topic_id, topic):
    payload = {
        "schema_version": LONG_TEXT_REVIEW_SCHEMA_VERSION,
        "topic_id": int(topic_id),
        "title": topic.get("title", ""),
        "fields": topic.get("fields", {}),
    }
    canonical = json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()[:24]


def _structure_report_text(report):
    issues = report["issues"]
    grouped = Counter(item["code"] for item in issues)
    lines = [
        "تقرير الفحص البنيوي لمخرجات تكست لونج",
        f"النتيجة: {'سليم' if report['structure_passed'] else 'غير سليم'}",
        f"الموضوعات المتوقعة: {report['topics_expected']}",
        f"الموضوعات المقروءة: {report['topics_read']}",
        f"إجمالي الأخطاء: {report['error_count']}",
        f"إجمالي التحذيرات: {report['warning_count']}",
        "",
    ]
    if grouped:
        lines.append("ملخص الأكواد:")
        for code, count in sorted(grouped.items()):
            lines.append(f"- {code}: {count}")
        lines.append("")
        lines.append("التفاصيل الكاملة:")
        for index, item in enumerate(issues, start=1):
            scope = []
            if item.get("topic_id") is not None:
                scope.append(f"موضوع {item['topic_id']}")
            if item.get("file"):
                scope.append(item["file"])
            if item.get("location"):
                scope.append(item["location"])
            scope_text = " — ".join(scope)
            detail = f"{index}. [{item['severity']}] {item['code']}"
            if scope_text:
                detail += f" — {scope_text}"
            detail += f": {item['message']}"
            lines.append(detail)
    else:
        lines.append("لا توجد مخالفات بنيوية أو تنسيقية.")
    return "\n".join(lines)


def _write_structure_reports(ctx, report, *, json_name, text_name, regenerate_name):
    with open(ctx.output_path(json_name), "w", encoding="utf-8") as handle:
        json.dump(report, handle, ensure_ascii=False, indent=2)
    with open(ctx.output_path(text_name), "w", encoding="utf-8") as handle:
        handle.write(_structure_report_text(report))

    regenerate = defaultdict(list)
    global_reasons = []
    for item in report["issues"]:
        if item["severity"] != "error":
            continue
        if item.get("topic_id") is None:
            global_reasons.append(
                {"code": item["code"], "message": item["message"]}
            )
        else:
            regenerate[str(item["topic_id"])].append(
                {"code": item["code"], "message": item["message"]}
            )
    payload = {
        "schema_version": LONG_TEXT_REVIEW_SCHEMA_VERSION,
        "global_blockers": global_reasons,
        "topics": [
            {
                "topic_id": int(topic_id),
                "reasons": reasons,
            }
            for topic_id, reasons in sorted(regenerate.items(), key=lambda pair: int(pair[0]))
        ],
    }
    with open(ctx.output_path(regenerate_name), "w", encoding="utf-8") as handle:
        json.dump(payload, handle, ensure_ascii=False, indent=2)


def action_long_text_review_build_cards(step, ctx):
    """Read final DOCX outputs and enforce the complete deterministic contract."""
    source_format = FORMATTING["source"]
    thumbnail_format = FORMATTING["thumbnail"]
    source_file = step.get("source_file", CONSUMER_EXACT["source_file"])
    thumbnail_file = step.get("thumbnail_file", CONSUMER_EXACT["thumbnail_file"])
    topics_file = step.get("topics_file", CONSUMER_EXACT["topics_file"])
    require_topic_ids = _bool(step.get("require_topic_ids", True), True)
    strict = _bool(step.get("strict", True), True)
    expected_font = step.get("font_name", source_format["font_name"])
    expected_size = float(step.get("font_size", source_format["font_size_pt"]))
    expected_line_spacing = float(
        step.get("line_spacing", source_format["line_spacing_pt"])
    )
    thumbnail_font = step.get(
        "thumbnail_font_name",
        thumbnail_format["font_name"],
    )
    thumbnail_size = float(
        step.get("thumbnail_font_size", thumbnail_format["font_size_pt"])
    )
    thumbnail_spacing = float(
        step.get(
            "thumbnail_line_spacing",
            thumbnail_format["line_spacing_pt"],
        )
    )

    issues = []
    seen = set()

    paths = {}
    for key, filename in (
        ("source", source_file),
        ("thumbnail", thumbnail_file),
        ("topics", topics_file),
    ):
        path = ctx.input_path(filename)
        paths[key] = path
        if not os.path.isfile(path):
            _add_unique(
                issues,
                _issue(
                    "REVIEW_INPUT_MISSING",
                    f"ملف الإدخال المطلوب غير موجود: {filename}",
                    file=filename,
                ),
                seen,
            )

    # Ambiguous fallback files are dangerous because a downstream recipe can
    # accidentally open a stale canonical file.
    try:
        source_stem = Path(source_file).stem
        matching = sorted(
            name for name in os.listdir(ctx.input_dir)
            if name.lower().endswith(".docx") and name.startswith(source_stem)
        )
        if len(matching) > 1:
            _add_unique(
                issues,
                _issue(
                    "AMBIGUOUS_SOURCE_OUTPUT_FILES",
                    "مجلد الإدخال يحتوي أكثر من نسخة محتملة لملف scripts_output",
                    file=source_file,
                    actual=matching,
                ),
                seen,
            )
        failures_path = ctx.input_path("generate_failures.json")
        if os.path.isfile(failures_path):
            try:
                failures = _strict_json_load(failures_path)
            except Exception as exc:
                failures = {"unreadable": str(exc)}
            _add_unique(
                issues,
                _issue(
                    "GENERATION_FAILURE_MANIFEST_PRESENT",
                    "يوجد generate_failures.json؛ التشغيل الأصلي احتوى فشلا جزئيا",
                    file="generate_failures.json",
                    actual=failures,
                ),
                seen,
            )
    except OSError as exc:
        _add_unique(
            issues,
            _issue(
                "INPUT_DIRECTORY_UNREADABLE",
                f"تعذر فحص مجلد الإدخال: {exc}",
            ),
            seen,
        )

    titles = {}
    topic_order = []
    if os.path.isfile(paths["topics"]):
        titles, topic_order, _metadata = _read_topics(paths["topics"], issues, seen)

    source_entries = []
    if os.path.isfile(paths["source"]):
        source_entries, _source_document = _load_document_topics(
            paths["source"],
            role=source_file,
            expected_font=expected_font,
            expected_size=expected_size,
            expected_line_spacing=expected_line_spacing,
            issues=issues,
            seen=seen,
        )

    if getattr(ctx, "topic_ids", None):
        expected_ids = []
        selected_seen = set()
        for index, value in enumerate(ctx.topic_ids, start=1):
            try:
                topic_id = _canonical_positive_id(value, f"TOPIC_IDS[{index}]")
            except ValueError as exc:
                _add_unique(
                    issues,
                    _issue(
                        "SELECTED_TOPIC_ID_INVALID",
                        str(exc),
                        location=f"TOPIC_IDS[{index}]",
                        actual=value,
                    ),
                    seen,
                )
                continue
            if topic_id in selected_seen:
                _add_unique(
                    issues,
                    _issue(
                        "SELECTED_TOPIC_ID_DUPLICATE",
                        f"المعرف {topic_id} مكرر في اختيار التشغيل",
                        topic_id=topic_id,
                        location=f"TOPIC_IDS[{index}]",
                    ),
                    seen,
                )
                continue
            selected_seen.add(topic_id)
            expected_ids.append(topic_id)
    elif require_topic_ids:
        expected_ids = []
        _add_unique(
            issues,
            _issue(
                "TOPIC_IDS_REQUIRED",
                "لا يمكن إثبات الموضوعات المفقودة بدون اختيار TOPIC_IDS نفسها المستخدمة في تشغيل التوليد",
                expected="اختيار معرفات التشغيل الأصلي من الواجهة",
            ),
            seen,
        )
    else:
        expected_ids = list(topic_order)

    if source_entries:
        _validate_id_coverage(
            source_entries,
            titles,
            expected_ids,
            source_file,
            issues,
            seen,
        )

    parsed_topics = {}
    entry_by_id = {}
    for entry in source_entries:
        topic_id = entry["id"]
        if topic_id in entry_by_id:
            continue
        entry_by_id[topic_id] = entry
        fields, sequence = _parse_script_fields(entry, source_file, issues, seen)
        title = titles.get(topic_id, "")
        parsed_topics[topic_id] = {
            "title": title,
            "fields": fields,
            "sequence": sequence,
        }
        _validate_field_content(
            fields,
            title,
            titles,
            topic_id,
            source_file,
            issues,
            seen,
        )

    # Detect exact copy/shift across topic and platform field names. Hashtags
    # are removed from descriptions first, so changing only tags cannot hide
    # a copied body. Similar/paraphrased transfers are handled by the global
    # semantic gate because local fuzzy matching would create false positives.
    duplicate_families = {
        "youtube_title_2": "title",
        "tiktok_title": "title",
        "facebook_title": "title",
        "youtube_description_1": "description",
        "youtube_description_2": "description",
        "tiktok_description": "description",
        "facebook_description": "description",
        "youtube_thumbnail_1": "short_phrase",
        "youtube_thumbnail_2": "short_phrase",
        "tiktok_screen": "short_phrase",
        "facebook_thumbnail": "short_phrase",
        "translation_en_title": "translation_en",
        "translation_en_description": "translation_en",
        "translation_fr_title": "translation_fr",
        "translation_fr_description": "translation_fr",
        "translation_es_title": "translation_es",
        "translation_es_description": "translation_es",
        "translation_de_title": "translation_de",
        "translation_de_description": "translation_de",
    }
    duplicates = defaultdict(list)
    for topic_id, topic in parsed_topics.items():
        for field_key, value in topic.get("fields", {}).items():
            family = duplicate_families.get(field_key)
            if not family:
                continue
            body = _strip_hashtags(value) if family == "description" else value
            normalized = _arabic_norm(body)
            if len(normalized) >= COUNTS["cross_topic_exact_duplicate_min_normalized_chars"]:
                duplicates[(family, normalized)].append((topic_id, field_key))
    for (family, _normalized), matches in duplicates.items():
        topic_ids = list(dict.fromkeys(topic_id for topic_id, _key in matches))
        if len(topic_ids) > 1:
            for topic_id, field_key in matches:
                _add_unique(
                    issues,
                    _issue(
                        "CROSS_TOPIC_FIELD_DUPLICATE",
                        "نفس المحتوى الجوهري مكرر حرفيا أو منقول بين "
                        f"الموضوعات {', '.join(topic_ids)}",
                        topic_id=topic_id,
                        file=source_file,
                        location=field_key,
                        actual={
                            "family": family,
                            "matches": [
                                {"topic_id": int(found_id), "field": found_key}
                                for found_id, found_key in matches
                            ],
                        },
                    ),
                    seen,
                )

    thumbnail_entries = []
    if os.path.isfile(paths["thumbnail"]):
        thumbnail_entries, _thumbnail_document = _load_document_topics(
            paths["thumbnail"],
            role=thumbnail_file,
            expected_font=thumbnail_font,
            expected_size=thumbnail_size,
            expected_line_spacing=thumbnail_spacing,
            issues=issues,
            seen=seen,
        )
        _validate_thumbnail_doc(
            thumbnail_entries,
            expected_ids,
            parsed_topics,
            thumbnail_file,
            issues,
            seen,
        )

    global_errors = [
        item for item in issues
        if item["severity"] == "error" and item.get("topic_id") is None
    ]
    blocked = defaultdict(list)
    for item in issues:
        if item["severity"] == "error" and item.get("topic_id") is not None:
            blocked[str(item["topic_id"])].append(
                f"{item['code']}: {item['message']}"
            )
    if global_errors:
        for topic_id in expected_ids:
            blocked[topic_id].extend(
                f"{item['code']}: {item['message']}" for item in global_errors
            )

    cards_topics = {}
    for topic_id in expected_ids:
        topic = parsed_topics.get(
            topic_id,
            {"title": titles.get(topic_id, ""), "fields": {}, "sequence": []},
        )
        topic["request_id"] = _request_id(topic_id, topic)
        cards_topics[topic_id] = topic

    error_count = sum(item["severity"] == "error" for item in issues)
    warning_count = sum(item["severity"] == "warning" for item in issues)
    report = {
        "schema_version": LONG_TEXT_REVIEW_SCHEMA_VERSION,
        "structure_passed": error_count == 0,
        "topics_expected": len(expected_ids),
        "topics_read": len({entry["id"] for entry in source_entries}),
        "expected_ids": [int(value) for value in expected_ids],
        "source_ids": [int(entry["id"]) for entry in source_entries],
        "thumbnail_ids": [int(entry["id"]) for entry in thumbnail_entries],
        "error_count": error_count,
        "warning_count": warning_count,
        "ingestion_usable": error_count == 0,
        "quality_clean": error_count == 0 and warning_count == 0,
        "issues": issues,
    }
    json_name = step.get("save_json", "long_text_structure_report.json")
    text_name = step.get("save_text", "long_text_structure_report.txt")
    regenerate_name = step.get(
        "save_regenerate", "long_text_topics_to_regenerate.json"
    )
    _write_structure_reports(
        ctx,
        report,
        json_name=json_name,
        text_name=text_name,
        regenerate_name=regenerate_name,
    )

    cards = {
        "schema_version": LONG_TEXT_REVIEW_SCHEMA_VERSION,
        "expected_ids": expected_ids,
        "topics": cards_topics,
        "blocked": {key: list(dict.fromkeys(value)) for key, value in blocked.items()},
        "issues": issues,
        "structure_report": report,
        "topic_filter": [int(value) for value in expected_ids]
        if getattr(ctx, "topic_ids", None)
        else None,
    }
    log(
        f"  long_text_review_build_cards: {len(cards_topics)} موضوع | "
        f"أخطاء={error_count} | تحذيرات={warning_count}"
    )
    if strict and error_count:
        raise EngineError(
            f"فشل الفحص البنيوي لمخرجات تكست لونج ({error_count} خطأ). "
            f"راجع {text_name} و{json_name}",
            code="LONG_TEXT_STRUCTURE_FAILED",
        )
    return cards

