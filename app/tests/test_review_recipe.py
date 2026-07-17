import copy
import json
import os
import sys
import tempfile
import unittest
from pathlib import Path
from unittest import mock

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
import recipe_runner as rr
from engine import BatchInfo, EngineResult


MARKER = "تفاصيل أكتر في الفيديو التوضيحي التالي"
ALT_MARKER = "لمزيد من التفاصيل شاهدوا الفيديو التوضيحي التالي"


class DummyContext:
    def __init__(self, input_dir, output_dir, topic_ids=None):
        self.input_dir = str(input_dir)
        self.output_dir = str(output_dir)
        self.topic_ids = topic_ids
        self.results = {}

    def input_path(self, filename):
        return os.path.join(self.input_dir, filename)

    def output_path(self, filename):
        return os.path.join(self.output_dir, filename)

    def resolve(self, value):
        if isinstance(value, str) and value.startswith("{") and value.endswith("}"):
            key = value[1:-1]
            if key in self.results:
                return self.results[key]
        return value

    def resolve_list(self, value):
        resolved = self.resolve(value)
        return resolved if isinstance(resolved, list) else [resolved]


class ReviewRecipeTests(unittest.TestCase):
    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.root = Path(self.temp_dir.name)
        self.input_dir = self.root / "input"
        self.output_dir = self.root / "output"
        self.input_dir.mkdir()
        self.output_dir.mkdir()
        self.ctx = DummyContext(self.input_dir, self.output_dir)

    def tearDown(self):
        self.temp_dir.cleanup()

    @staticmethod
    def long_intro(topic_id, pair, marker=MARKER):
        return (
            f"مقدمة {pair} عن أعراض مرض الكلى للموضوع {topic_id} تشرح العلامات المبكرة بوضوح "
            "وتربط تغير البول والتورم والإرهاق بوظيفة الترشيح حتى يفهم المشاهد محور الأعراض "
            "ويعرف لماذا تظهر العلامات بالتدريج ومتى يحتاج إلى تقييم طبي دقيق "
            f"{marker}"
        )

    @staticmethod
    def long_text(topic_id, pair):
        return (
            f"النص {pair} عن أعراض مرض الكلى للموضوع {topic_id} يوضح أن تغير البول والتورم والإرهاق "
            "قد ينتج عن تراجع الترشيح واحتباس السوائل وتراكم الفضلات في الجسم ويشرح العلاقة العلمية "
            "بين كل علامة ووظيفة الكلى مع تنبيه واضح إلى أهمية الفحص الطبي عند استمرار الأعراض"
        )

    def write_inputs(
        self,
        doc_topic_ids=(1,),
        title_items=None,
        intro_marker=MARKER,
        intro_mutator=None,
        text_mutator=None,
        bom=False,
    ):
        intro_doc = Document()
        text_doc = Document()
        for topic_id in doc_topic_ids:
            intro_doc.add_paragraph(f"Script {topic_id}")
            for pair in range(1, 5):
                text = self.long_intro(topic_id, pair, intro_marker)
                if intro_mutator:
                    text = intro_mutator(intro_doc, topic_id, pair, text)
                if text is not None:
                    intro_doc.add_paragraph(text)

            text_doc.add_paragraph(f"Script {topic_id}")
            for pair in range(1, 5):
                if text_mutator:
                    text_mutator(text_doc, topic_id, pair)
                text_doc.add_paragraph(f"Part {pair}")
                text_doc.add_paragraph(self.long_text(topic_id, pair))

        intro_doc.save(self.input_dir / "intros_output.docx")
        text_doc.save(self.input_dir / "texts_output.docx")
        if title_items is None:
            title_items = [
                {"id": topic_id, "title": f"أعراض مرض الكلى للموضوع {topic_id}"}
                for topic_id in doc_topic_ids
            ]
        encoding = "utf-8-sig" if bom else "utf-8"
        with open(self.input_dir / "topics.json", "w", encoding=encoding) as handle:
            json.dump({"titles": title_items}, handle, ensure_ascii=False)
        with open(self.input_dir / "instructions.txt", "w", encoding="utf-8") as handle:
            handle.write("اقرأ كل عنوان وكل مقدمة وكل نص، واحكم على المرض والمحور والمعنى بدقة كاملة.")

    def build_cards(self, strict=True, **overrides):
        step = {
            "id": "cards",
            "action": "review_build_cards",
            "intros_file": "intros_output.docx",
            "texts_file": "texts_output.docx",
            "topics_file": "topics.json",
            "expected_pairs": 4,
            "min_words": 1,
            "strict": strict,
        }
        step.update(overrides)
        return rr.action_review_build_cards(step, self.ctx)

    def build_prompts(self, cards, instructions=None):
        self.ctx.results["cards"] = cards
        self.ctx.results["instructions"] = (
            instructions
            if instructions is not None
            else "راجع كل زوج دلاليا ولا تعتمد على تشابه الكلمات فقط."
        )
        return rr.action_review_build_prompts(
            {
                "id": "prompts",
                "input": "{cards}",
                "instructions": "{instructions}",
                "save_as": "review_requests.json",
            },
            self.ctx,
        )

    @staticmethod
    def evidence_from(source):
        return " ".join(source.split()[:5])

    def valid_response(self, cards, topic_id="1"):
        card = cards["topics"][topic_id]
        pairs = []
        for pair in range(1, cards["expected_pairs"] + 1):
            pairs.append(
                {
                    "pair": pair,
                    "intro_focus": "أعراض مرض الكلى",
                    "text_focus": "أعراض مرض الكلى",
                    "intro_evidence": self.evidence_from(card["intros"][pair - 1]),
                    "text_evidence": self.evidence_from(card["texts"][str(pair)]),
                    "text_vs_intro": "مطابق",
                    "intro_vs_title": "مطابق",
                    "text_vs_title": "مطابق",
                    "reason": "",
                }
            )
        return {
            "schema_version": rr._REVIEW_SCHEMA_VERSION,
            "request_id": card["request_id"],
            "topic_id": int(topic_id),
            "pairs": pairs,
        }

    def parse_responses(self, cards, responses):
        self.ctx.results["cards"] = cards
        self.ctx.results["responses"] = responses
        report = rr.action_review_parse_verdicts(
            {
                "id": "report",
                "input": "{responses}",
                "cards": "{cards}",
                "save_json": "review_report.json",
            },
            self.ctx,
        )
        with open(self.output_dir / "review_report.json", encoding="utf-8") as handle:
            data = json.load(handle)
        return report, data

    def test_valid_files_build_clean_card(self):
        self.write_inputs()
        cards = self.build_cards()
        self.assertEqual([], cards["issues"])
        self.assertEqual(["1"], list(cards["topics"]))
        self.assertEqual(4, len(cards["topics"]["1"]["intros"]))
        self.assertRegex(cards["topics"]["1"]["request_id"], r"^[0-9a-f]{20}$")

    def test_alternative_intro_marker_is_supported(self):
        self.write_inputs(intro_marker=ALT_MARKER)
        cards = self.build_cards()
        self.assertEqual([], cards["issues"])

    def test_title_only_topic_is_reported_missing(self):
        self.write_inputs(
            doc_topic_ids=(1,),
            title_items=[{"id": 1, "title": "عنوان 1"}, {"id": 2, "title": "عنوان 2"}],
        )
        cards = self.build_cards(strict=False)
        self.assertIn("2", cards["topics"])
        self.assertIn("2", cards["blocked"])
        self.assertTrue(any("المقدمات مفقودة بالكامل" in item for item in cards["blocked"]["2"]))
        self.assertTrue(any("النصوص مفقودة بالكامل" in item for item in cards["blocked"]["2"]))

    def test_selected_unknown_topic_is_not_silently_dropped(self):
        self.write_inputs()
        self.ctx.topic_ids = {1, 2}
        cards = self.build_cards(strict=False)
        self.assertEqual({"1", "2"}, set(cards["topics"]))
        self.assertIn("2", cards["blocked"])

    def test_duplicate_title_id_is_blocked(self):
        self.write_inputs(
            title_items=[
                {"id": 1, "title": "العنوان الأول"},
                {"id": 1, "title": "العنوان الثاني"},
            ]
        )
        cards = self.build_cards(strict=False)
        self.assertIn("1", cards["blocked"])
        self.assertEqual("العنوان الأول", cards["topics"]["1"]["title"])
        self.assertTrue(any("مكرر" in item for item in cards["blocked"]["1"]))

    def test_fractional_title_id_is_rejected(self):
        self.write_inputs(title_items=[{"id": 1.9, "title": "عنوان"}])
        with self.assertRaises(rr.EngineError) as raised:
            self.build_cards(strict=False)
        self.assertEqual("REVIEW_TOPICS_EMPTY", raised.exception.code)
        self.assertTrue((self.output_dir / "structure_report.txt").exists())

    def test_boolean_title_id_is_rejected(self):
        self.write_inputs(title_items=[{"id": True, "title": "عنوان"}])
        with self.assertRaises(rr.EngineError):
            self.build_cards(strict=False)

    def test_null_title_is_blocked_not_stringified(self):
        self.write_inputs(title_items=[{"id": 1, "title": None}])
        cards = self.build_cards(strict=False)
        self.assertEqual("", cards["topics"]["1"]["title"])
        self.assertIn("1", cards["blocked"])

    def test_utf8_bom_topics_file_is_supported(self):
        self.write_inputs(bom=True)
        cards = self.build_cards()
        self.assertEqual([], cards["issues"])

    def test_malformed_topics_json_writes_structure_report(self):
        self.write_inputs()
        (self.input_dir / "topics.json").write_text("{broken", encoding="utf-8")
        with self.assertRaises(rr.EngineError) as raised:
            self.build_cards()
        self.assertEqual("REVIEW_TOPICS_INVALID", raised.exception.code)
        self.assertIn(
            "فشل فحص المدخلات",
            (self.output_dir / "structure_report.txt").read_text(encoding="utf-8"),
        )

    def test_last_intro_without_marker_is_blocked(self):
        def mutate(_doc, _topic, pair, text):
            return text.replace(MARKER, "") if pair == 4 else text

        self.write_inputs(intro_mutator=mutate)
        cards = self.build_cards(strict=False)
        self.assertIn("1", cards["blocked"])
        self.assertTrue(any("ماركرات" in item for item in cards["blocked"]["1"]))

    def test_four_markerless_paragraphs_do_not_bypass_gate(self):
        self.write_inputs(intro_marker="")
        cards = self.build_cards(strict=False)
        self.assertIn("1", cards["blocked"])
        self.assertTrue(any("عدد ماركرات" in item for item in cards["blocked"]["1"]))

    def test_orphan_text_between_script_and_first_part_is_blocked(self):
        def mutate(document, _topic, pair):
            if pair == 1:
                document.add_paragraph("هذا نص يتيم لا يحمل عنوان Part ولا يجوز إسقاطه")

        self.write_inputs(text_mutator=mutate)
        cards = self.build_cards(strict=False)
        self.assertIn("1", cards["blocked"])
        self.assertTrue(any("محتوى يتيم" in item for item in cards["blocked"]["1"]))

    def test_near_miss_script_header_is_reported(self):
        self.write_inputs()
        bad_doc = Document()
        bad_doc.add_paragraph("Script 1:")
        for pair in range(1, 5):
            bad_doc.add_paragraph(self.long_intro(1, pair))
        bad_doc.save(self.input_dir / "intros_output.docx")
        cards = self.build_cards(strict=False)
        self.assertTrue(any("فاصل Script غير صالح" in issue for issue in cards["issues"]))
        self.assertIn("1", cards["blocked"])

    def test_huge_id_gap_is_counted_without_expanding_range(self):
        count, sample = rr._review_sequence_gaps(["1", "1000000000"])
        self.assertEqual(999999998, count)
        self.assertEqual(30, len(sample))
        self.assertEqual("2", sample[0])

    def test_bad_expected_pairs_type_fails_preflight(self):
        self.write_inputs()
        with self.assertRaises(rr.EngineError) as raised:
            self.build_cards(expected_pairs=4.9)
        self.assertEqual("REVIEW_BAD_PARAMS", raised.exception.code)

    def test_bad_min_words_type_fails_preflight(self):
        self.write_inputs()
        with self.assertRaises(rr.EngineError) as raised:
            self.build_cards(min_words=1.5)
        self.assertEqual("REVIEW_BAD_PARAMS", raised.exception.code)

    def test_empty_instructions_stop_before_batch(self):
        self.write_inputs()
        cards = self.build_cards()
        with self.assertRaises(rr.EngineError) as raised:
            self.build_prompts(cards, instructions="")
        self.assertEqual("REVIEW_INSTRUCTIONS_EMPTY", raised.exception.code)

    def test_prompt_contains_one_json_card_and_audit_file(self):
        self.write_inputs(doc_topic_ids=(1, 2))
        cards = self.build_cards()
        prompts = self.build_prompts(cards)
        self.assertEqual(2, len(prompts))
        for prompt, topic_id in zip(prompts, (1, 2)):
            start = prompt.index("<BEGIN_UNTRUSTED_REVIEW_DATA>") + len("<BEGIN_UNTRUSTED_REVIEW_DATA>")
            end = prompt.index("<END_UNTRUSTED_REVIEW_DATA>")
            payload = json.loads(prompt[start:end].strip())
            self.assertEqual(topic_id, payload["topic_id"])
            self.assertEqual(4, len(payload["pairs"]))
            self.assertGreater(prompt.rfind("تنبيه أمان نهائي"), end)
        self.assertTrue((self.output_dir / "review_requests.json").exists())

    def test_valid_response_with_verified_evidence_succeeds(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        report, data = self.parse_responses(cards, response)
        self.assertIn("✅ كل الموضوعات مطابقة", report)
        self.assertEqual(1, data["judged"])
        self.assertEqual([], data["unjudged"])

    def test_fractional_topic_id_cannot_alias_integer(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["topic_id"] = 1.9
        report, data = self.parse_responses(cards, response)
        self.assertNotIn("✅ كل الموضوعات مطابقة", report)
        self.assertEqual(0, data["judged"])
        self.assertEqual(1, len(data["unparseable_responses"]))

    def test_fractional_pair_cannot_alias_integer(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][0]["pair"] = 1.9
        report, data = self.parse_responses(cards, response)
        self.assertNotIn("✅ كل الموضوعات مطابقة", report)
        self.assertEqual(0, data["judged"])
        self.assertTrue(data["unjudged"])

    def test_extra_pair_blocks_success(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        extra = copy.deepcopy(response["pairs"][-1])
        extra["pair"] = 5
        response["pairs"].append(extra)
        report, data = self.parse_responses(cards, response)
        self.assertNotIn("✅ كل الموضوعات مطابقة", report)
        self.assertEqual(0, data["judged"])

    def test_empty_pairs_is_not_counted_as_judged(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"] = []
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])
        self.assertEqual("غير محكوم جزئياً", data["detail"]["1"]["status"])

    def test_duplicate_pair_blocks_success(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][1]["pair"] = 1
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])
        self.assertTrue(any("مكرر" in item["reason"] for item in data["unjudged"]))

    def test_wrong_request_id_blocks_success(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["request_id"] = "0" * 20
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])
        self.assertTrue(data["response_schema_errors"])

    def test_missing_evidence_blocks_success(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][0]["intro_evidence"] = ""
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])
        self.assertTrue(any("intro_evidence" in item["reason"] for item in data["unjudged"]))

    def test_hallucinated_evidence_blocks_success(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][0]["text_evidence"] = "كلام غير موجود داخل النص إطلاقا"
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])
        self.assertTrue(any("مش موجود حرفياً" in item["reason"] for item in data["unjudged"]))

    def test_extra_pair_field_blocks_success(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][0]["unexpected"] = "value"
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])
        self.assertTrue(any("حقول زائدة" in item["reason"] for item in data["unjudged"]))

    def test_extra_top_level_field_blocks_success(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["summary"] = "مطابق"
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])
        self.assertTrue(data["response_schema_errors"])

    def test_mismatch_without_reason_is_also_unjudged(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][0]["text_vs_intro"] = "غير مطابق"
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, len(data["mismatches"]))
        self.assertEqual(0, data["judged"])
        self.assertTrue(any("reason مطلوب" in item["reason"] for item in data["unjudged"]))

    def test_valid_mismatch_is_reported_without_becoming_unjudged(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][0]["text_vs_intro"] = "غير مطابق"
        response["pairs"][0]["reason"] = "النص انتقل إلى محور العلاج بدل شرح الأعراض"
        report, data = self.parse_responses(cards, response)
        self.assertIn("غير مطابق", report)
        self.assertEqual(1, data["judged"])
        self.assertEqual([], data["unjudged"])
        self.assertEqual(1, len(data["mismatches"]))

    def test_reason_with_all_matching_is_rejected(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][0]["reason"] = "كل شيء مطابق"
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])
        self.assertTrue(any("لازم يكون فارغاً" in item["reason"] for item in data["unjudged"]))

    def test_non_string_reason_on_mismatch_is_rejected(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][0]["text_vs_intro"] = "غير مطابق"
        response["pairs"][0]["reason"] = None
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])
        self.assertTrue(any("reason لازم يكون نصاً" in item["reason"] for item in data["unjudged"]))

    def test_duplicate_topic_response_is_unjudged(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        _report, data = self.parse_responses(cards, [response, copy.deepcopy(response)])
        self.assertEqual(0, data["judged"])
        self.assertEqual(["1"], data["duplicate_topic_responses"])

    def test_alien_response_prevents_clean_success(self):
        self.write_inputs()
        cards = self.build_cards()
        valid = self.valid_response(cards)
        alien = copy.deepcopy(valid)
        alien["topic_id"] = 999
        report, data = self.parse_responses(cards, [valid, alien])
        self.assertNotIn("✅ كل الموضوعات مطابقة", report)
        self.assertEqual(1, data["judged"])
        self.assertEqual(1, len(data["alien_responses"]))

    def test_unparseable_extra_response_prevents_clean_success(self):
        self.write_inputs()
        cards = self.build_cards()
        valid = self.valid_response(cards)
        report, data = self.parse_responses(cards, [valid, "not json"])
        self.assertNotIn("✅ كل الموضوعات مطابقة", report)
        self.assertEqual(1, len(data["unparseable_responses"]))

    def test_raw_dict_response_is_supported(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        report, data = self.parse_responses(cards, response)
        self.assertIn("✅ كل الموضوعات مطابقة", report)
        self.assertEqual(1, data["judged"])

    def test_no_topics_can_never_return_success(self):
        cards = {
            "schema_version": rr._REVIEW_SCHEMA_VERSION,
            "expected_pairs": 4,
            "topics": {},
            "blocked": {},
            "issues": [],
        }
        with self.assertRaises(rr.EngineError) as raised:
            self.parse_responses(cards, [])
        self.assertEqual("REVIEW_NO_TOPICS", raised.exception.code)

    def test_all_blocked_topics_are_not_sent(self):
        self.write_inputs(
            doc_topic_ids=(1,),
            title_items=[{"id": 1, "title": "عنوان 1"}, {"id": 2, "title": "عنوان 2"}],
        )
        self.ctx.topic_ids = {2}
        cards = self.build_cards(strict=False)
        with self.assertRaises(rr.EngineError) as raised:
            self.build_prompts(cards)
        self.assertEqual("REVIEW_ALL_BLOCKED", raised.exception.code)

    def test_exact_duplicate_content_is_reported(self):
        self.write_inputs()
        text_doc = Document()
        text_doc.add_paragraph("Script 1")
        duplicate = self.long_text(1, 1)
        for pair in range(1, 5):
            text_doc.add_paragraph(f"Part {pair}")
            text_doc.add_paragraph(duplicate)
        text_doc.save(self.input_dir / "texts_output.docx")
        cards = self.build_cards(strict=False)
        self.assertTrue(any("تكرار حرفي" in issue for issue in cards["issues"]))


    def test_placeholder_focus_is_rejected(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][0]["intro_focus"] = "x"
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])

    def test_shared_generic_evidence_is_rejected(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        shared = "\u0639\u0646 \u0623\u0639\u0631\u0627\u0636 \u0645\u0631\u0636 \u0627\u0644\u0643\u0644\u0649"
        response["pairs"][0]["intro_evidence"] = shared
        response["pairs"][0]["text_evidence"] = shared
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])

    def test_stale_card_fingerprint_is_rejected(self):
        self.write_inputs()
        cards = self.build_cards()
        cards["topics"]["1"]["texts"]["1"] += " changed"
        response = self.valid_response(cards)
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])
        self.assertTrue(data["response_schema_errors"])

    def test_chatter_around_json_is_rejected(self):
        self.write_inputs()
        cards = self.build_cards()
        raw = "before " + json.dumps(self.valid_response(cards), ensure_ascii=False)
        _report, data = self.parse_responses(cards, raw)
        self.assertEqual(0, data["judged"])
        self.assertTrue(data["unparseable_responses"])

    def test_duplicate_json_member_is_rejected(self):
        self.write_inputs()
        cards = self.build_cards()
        raw = json.dumps(self.valid_response(cards), ensure_ascii=False)
        raw = raw.replace("{", '{"topic_id": 1,', 1)
        _report, data = self.parse_responses(cards, raw)
        self.assertEqual(0, data["judged"])
        self.assertTrue(data["unparseable_responses"])

    def test_numeric_strings_in_response_identity_are_rejected(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["schema_version"] = str(response["schema_version"])
        response["topic_id"] = str(response["topic_id"])
        response["pairs"][0]["pair"] = "1"
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])

    def test_overlong_enum_cannot_pass_by_truncation(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][0]["text_vs_intro"] += "\u064e" * 75 + "TRAILING_INVALID"
        _report, data = self.parse_responses(cards, response)
        self.assertEqual(0, data["judged"])

    def test_schema_invalid_pair_does_not_publish_mismatch(self):
        self.write_inputs()
        cards = self.build_cards()
        response = self.valid_response(cards)
        response["pairs"][0]["unexpected"] = True
        response["pairs"][0]["text_vs_intro"] = "\u063a\u064a\u0631 \u0645\u0637\u0627\u0628\u0642"
        response["pairs"][0]["reason"] = "axis mismatch"
        _report, data = self.parse_responses(cards, response)
        self.assertEqual([], data["mismatches"])
        self.assertEqual(0, data["judged"])

    def test_script_zero_becomes_documented_structure_error(self):
        self.write_inputs()
        document = Document(self.input_dir / "intros_output.docx")
        document.paragraphs[0].text = "Script 0"
        document.save(self.input_dir / "intros_output.docx")
        with self.assertRaises(rr.EngineError) as raised:
            self.build_cards(strict=False)
        self.assertEqual("REVIEW_STRUCTURE_FAILED", raised.exception.code)
        self.assertTrue((self.output_dir / "structure_report.txt").exists())

    def test_ambiguous_titles_and_topics_keys_are_rejected(self):
        self.write_inputs()
        payload = {
            "titles": [{"id": 1, "title": "title one"}],
            "topics": [{"id": 2, "title": "hidden title"}],
        }
        with open(self.input_dir / "topics.json", "w", encoding="utf-8") as handle:
            json.dump(payload, handle)
        with self.assertRaises(rr.EngineError) as raised:
            self.build_cards(strict=False)
        self.assertEqual("REVIEW_TOPICS_AMBIGUOUS", raised.exception.code)

    def test_word_table_content_is_never_ignored(self):
        self.write_inputs()
        document = Document(self.input_dir / "intros_output.docx")
        document.add_table(rows=1, cols=1).cell(0, 0).text = "Script 2 hidden content"
        document.save(self.input_dir / "intros_output.docx")
        cards = self.build_cards(strict=False)
        self.assertTrue(any("Word" in issue for issue in cards["issues"]))

    def test_two_markers_in_one_intro_paragraph_are_blocked(self):
        self.write_inputs()
        document = Document(self.input_dir / "intros_output.docx")
        document.paragraphs[1].text += " " + MARKER
        document.save(self.input_dir / "intros_output.docx")
        cards = self.build_cards(strict=False)
        self.assertIn("1", cards["blocked"])

    def test_contiguous_id_requirement_blocks_hidden_gap(self):
        self.write_inputs(doc_topic_ids=(1, 3))
        cards = self.build_cards(strict=False, require_contiguous_ids=True)
        self.assertTrue(any("(1)" in issue and "2" in issue for issue in cards["issues"]))

    def test_full_scope_recipe_forbids_topic_filter(self):
        self.write_inputs()
        self.ctx.topic_ids = {1}
        with self.assertRaises(rr.EngineError) as raised:
            self.build_cards(strict=False, allow_topic_filter=False)
        self.assertEqual("REVIEW_TOPIC_FILTER_FORBIDDEN", raised.exception.code)

    def test_invalid_topic_ids_environment_is_rejected(self):
        env = {
            "TOPIC_IDS": "1,abc",
            "INPUT_DIR": str(self.input_dir),
            "OUTPUT_DIR": str(self.output_dir),
        }
        with mock.patch.dict(os.environ, env, clear=False):
            with self.assertRaises(rr.EngineError) as raised:
                rr.PipelineContext()
        self.assertEqual("INVALID_TOPIC_IDS", raised.exception.code)

    def test_incomplete_batch_result_saves_report_then_fails(self):
        self.write_inputs(doc_topic_ids=(1, 2))
        cards = self.build_cards()
        self.ctx.results["cards"] = cards
        self.ctx.results["responses"] = [self.valid_response(cards, "1")]
        with self.assertRaises(rr.EngineError) as raised:
            rr.action_review_parse_verdicts(
                {
                    "id": "report",
                    "input": "{responses}",
                    "cards": "{cards}",
                    "save_json": "review_report.json",
                    "save_text": "review_report.txt",
                    "fail_incomplete": True,
                },
                self.ctx,
            )
        self.assertEqual("REVIEW_INCOMPLETE", raised.exception.code)
        self.assertTrue((self.output_dir / "review_report.json").exists())
        self.assertTrue((self.output_dir / "review_report.txt").exists())
    def load_review_recipe(self):
        recipes_dir = Path(rr.__file__).resolve().parents[1] / "recipes"
        for candidate in recipes_dir.glob("*.json"):
            with open(candidate, encoding="utf-8") as handle:
                config = json.load(handle)
            if any(
                step.get("action") == "review_build_cards"
                for step in config.get("steps", [])
            ):
                return config
        self.fail("review recipe not found")

    def test_explicit_batch_send_only_never_retrieves(self):
        self.write_inputs()
        config = self.load_review_recipe()

        def fake_send(**kwargs):
            info = BatchInfo(
                provider="vertex", model=kwargs["model"], job_id="send-only-job"
            )
            info.save(kwargs["save_path"])
            return EngineResult(success=True, data=info, provider="vertex")

        env = {
            "INPUT_DIR": str(self.input_dir),
            "RECIPE_OUTPUT_DIR": str(self.output_dir),
            "OUTPUT_DIR": str(self.output_dir),
            "EXECUTION_MODE": "send_only",
            "MODEL_NAME": "gemini-test",
            "TOPIC_IDS": "",
        }
        with mock.patch.dict(os.environ, env, clear=False), mock.patch.object(
            rr, "batch_send", side_effect=fake_send
        ), mock.patch.object(
            rr, "batch_retrieve", side_effect=AssertionError("retrieve called")
        ):
            rr.run_pipeline(config)
        self.assertTrue((self.output_dir / "batch_job_info.json").exists())
        self.assertFalse((self.output_dir / "review_report.json").exists())

    def test_explicit_batch_receive_only_never_sends(self):
        self.write_inputs()
        config = self.load_review_recipe()
        cards = self.build_cards()
        response = self.valid_response(cards)
        BatchInfo(
            provider="vertex", model="gemini-test", job_id="receive-only-job"
        ).save(self.output_dir / "batch_job_info.json")

        env = {
            "INPUT_DIR": str(self.input_dir),
            "RECIPE_OUTPUT_DIR": str(self.output_dir),
            "OUTPUT_DIR": str(self.output_dir),
            "EXECUTION_MODE": "receive_only",
            "MODEL_NAME": "gemini-test",
            "TOPIC_IDS": "",
        }
        with mock.patch.dict(os.environ, env, clear=False), mock.patch.object(
            rr, "batch_send", side_effect=AssertionError("send called")
        ), mock.patch.object(
            rr,
            "batch_retrieve",
            return_value=EngineResult(
                success=True,
                data=[json.dumps(response, ensure_ascii=False)],
                provider="vertex",
            ),
        ):
            rr.run_pipeline(config)
        data = json.loads(
            (self.output_dir / "review_report.json").read_text(encoding="utf-8")
        )
        self.assertTrue(data["review_completed"])

    def test_disallowed_batch_provider_stops_before_api(self):
        self.ctx.model = "gpt-5"
        self.ctx.thinking_level = "none"
        self.ctx.run_id = ""
        self.ctx.recipe_name = ""
        self.ctx.channel_name = ""
        with mock.patch.object(rr, "batch_send", side_effect=AssertionError("API called")):
            with self.assertRaises(rr.EngineError) as raised:
                rr.action_batch_send(
                    {
                        "id": "batch",
                        "prompts": ["prompt"],
                        "allowed_providers": ["gemini", "claude", "glm"],
                    },
                    self.ctx,
                )
        self.assertEqual("BATCH_PROVIDER_NOT_ALLOWED", raised.exception.code)
    def test_full_recipe_pipeline_with_mocked_batch_api(self):
        self.write_inputs()
        recipes_dir = Path(rr.__file__).resolve().parents[1] / "recipes"
        review_recipes = []
        for candidate in recipes_dir.glob("*.json"):
            with open(candidate, encoding="utf-8") as candidate_file:
                candidate_config = json.load(candidate_file)
            if any(
                step.get("action") == "review_build_cards"
                for step in candidate_config.get("steps", [])
            ):
                review_recipes.append(candidate)
        self.assertEqual(1, len(review_recipes))
        with open(review_recipes[0], encoding="utf-8") as handle:
            config = json.load(handle)
        self.assertEqual([], rr.validate_pipeline(config))

        captured = {}

        def fake_batch_send(**kwargs):
            captured["prompts"] = kwargs["prompts"]
            responses = []
            for prompt in kwargs["prompts"]:
                payload_text = prompt.split(
                    "<BEGIN_UNTRUSTED_REVIEW_DATA>\n", 1
                )[1].split("\n<END_UNTRUSTED_REVIEW_DATA>", 1)[0]
                payload = json.loads(payload_text)
                pairs = []
                for source_pair in payload["pairs"]:
                    pairs.append(
                        {
                            "pair": source_pair["pair"],
                            "intro_focus": "\u0623\u0639\u0631\u0627\u0636 \u0645\u0631\u0636 \u0627\u0644\u0643\u0644\u0649",
                            "text_focus": "\u0623\u0639\u0631\u0627\u0636 \u0645\u0631\u0636 \u0627\u0644\u0643\u0644\u0649",
                            "intro_evidence": self.evidence_from(source_pair["intro"]),
                            "text_evidence": self.evidence_from(source_pair["text"]),
                            "text_vs_intro": "\u0645\u0637\u0627\u0628\u0642",
                            "intro_vs_title": "\u0645\u0637\u0627\u0628\u0642",
                            "text_vs_title": "\u0645\u0637\u0627\u0628\u0642",
                            "reason": "",
                        }
                    )
                responses.append(
                    json.dumps(
                        {
                            "schema_version": payload["schema_version"],
                            "request_id": payload["request_id"],
                            "topic_id": payload["topic_id"],
                            "pairs": pairs,
                        },
                        ensure_ascii=False,
                    )
                )
            captured["responses"] = responses
            info = BatchInfo(
                provider="vertex",
                model=kwargs["model"],
                job_id="mock-review-job",
                items_count=len(responses),
                status="submitted",
            )
            info.save(kwargs["save_path"])
            return EngineResult(
                success=True,
                data=info,
                model=kwargs["model"],
                provider="vertex",
            )

        def fake_batch_retrieve(**_kwargs):
            return EngineResult(
                success=True,
                data=captured["responses"],
                model="gemini-test",
                provider="vertex",
            )

        env = {
            "INPUT_DIR": str(self.input_dir),
            "RECIPE_OUTPUT_DIR": str(self.output_dir),
            "OUTPUT_DIR": str(self.output_dir),
            "EXECUTION_MODE": "batch_auto",
            "MODEL_NAME": "gemini-test",
            "TOPIC_IDS": "",
            "RUN_ID": "review-e2e-test",
            "RECIPE_NAME": config["name"],
            "CHANNEL_NAME": "tests",
        }
        with mock.patch.dict(os.environ, env, clear=False), mock.patch.object(
            rr, "batch_send", side_effect=fake_batch_send
        ), mock.patch.object(
            rr, "batch_retrieve", side_effect=fake_batch_retrieve
        ):
            rr.run_pipeline(config)

        self.assertEqual(1, len(captured["prompts"]))
        self.assertTrue((self.output_dir / "batch_job_info.json").exists())
        self.assertTrue((self.output_dir / "review_requests.json").exists())
        self.assertTrue((self.output_dir / "review_responses_text.json").exists())
        self.assertTrue((self.output_dir / "review_report.json").exists())
        self.assertTrue((self.output_dir / "review_report.txt").exists())
        report = (self.output_dir / "review_report.txt").read_text(encoding="utf-8")
        self.assertTrue(report.strip())
        data = json.loads(
            (self.output_dir / "review_report.json").read_text(encoding="utf-8")
        )
        self.assertEqual(1, data["judged"])
        self.assertEqual(4, len(data["detail"]["1"]["pairs"]))
        self.assertEqual([], data["unjudged"])
if __name__ == "__main__":
    unittest.main(verbosity=2)
