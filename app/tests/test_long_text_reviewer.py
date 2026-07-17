import json
import os
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


APP_DIR = Path(__file__).resolve().parents[1]
if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

from long_text_reviewer import (  # noqa: E402
    CANONICAL_LABELS,
    LONG_TEXT_REVIEW_SCHEMA_VERSION,
    EXPECTED_TOKEN_KEYS,
    FIELD_KEYS,
    SEMANTIC_CROSS_CHECKS,
    TOKEN_KINDS,
    _parse_script_fields,
    _request_id,
    action_long_text_review_build_cards,
)
from long_text_semantic_reviewer import (  # noqa: E402
    _policy_sha256,
    _topic_review_request_id,
    action_long_text_review_build_prompts,
    action_long_text_review_parse_verdicts,
)


TITLE = "علامات مبكرة تستحق الانتباه لصحة الكلى"
REQUIRED_TAGS = "#صحة #الهشاب2000 #نصائح_طبية #ناتوريا"
KEYWORDS = ", ".join(f"كلمة مفتاحية {number}" for number in range(1, 19))


def valid_fields():
    return {
        "youtube_title_1": TITLE,
        "youtube_title_2": "مؤشرات أولية مهمة تساعد على متابعة صحة الكلى",
        "youtube_description_1": (
            "شرح علمي مبسط لعلامات قد تستدعي متابعة وظائف الكلى مع الطبيب. "
            f"{REQUIRED_TAGS} #صحة_الكلى #فحوصات"
        ),
        "youtube_description_2": (
            "توضيح عملي للمؤشرات المبكرة التي قد ترتبط بصحة الكلى وأهمية التقييم الطبي. "
            f"{REQUIRED_TAGS} #الكلى #الوقاية"
        ),
        "youtube_thumbnail_1": "علامات الكلى المهمة\nقبل ظهور التعب",
        "youtube_thumbnail_2": "صحة الكلى تبدأ\nبفهم هذه العلامات",
        "youtube_keywords": KEYWORDS,
        "tiktok_title": "علامات مهمة لصحة الكلى",
        "tiktok_description": (
            "علامات مهمة لصحة الكلى قد تساعدك على طلب تقييم طبي في الوقت المناسب. "
            f"{REQUIRED_TAGS} #الكلى #متابعة"
        ),
        "tiktok_screen": "افهم علامات الكلى\nقبل زيادة التعب",
        "facebook_title": "مؤشرات مبكرة لصحة الكلى",
        "facebook_description": (
            "مؤشرات مبكرة لصحة الكلى تستحق الانتباه والتقييم الطبي عند استمرارها. "
            f"{REQUIRED_TAGS} #فحص_الكلى #وعي"
        ),
        "facebook_thumbnail": "مؤشرات الكلى المهمة\nتستحق انتباهك",
        "facebook_keywords": KEYWORDS.replace("كلمة", "موضوع"),
        "translation_en_title": "Early signs worth noticing for kidney health",
        "translation_en_description": (
            "A clear guide to early signs that may justify medical kidney assessment. "
            "#kidney #health #medical #care #screening #awareness"
        ),
        "translation_fr_title": "Signes précoces à surveiller pour la santé rénale",
        "translation_fr_description": (
            "Un guide clair des signes pouvant justifier une évaluation médicale des reins. "
            "#reins #sante #medical #soins #depistage #prevention"
        ),
        "translation_es_title": "Señales tempranas para cuidar la salud renal",
        "translation_es_description": (
            "Una guía clara de señales que pueden justificar una evaluación médica renal. "
            "#rinon #salud #medicina #cuidados #revision #prevencion"
        ),
        "translation_de_title": "Frühe Hinweise für die Nierengesundheit",
        "translation_de_description": (
            "Ein klarer Überblick über Hinweise, die eine ärztliche Nierenprüfung begründen können. "
            "#niere #gesundheit #medizin #vorsorge #kontrolle #aufklaerung"
        ),
    }


def token_lines(fields=None):
    fields = fields or valid_fields()
    lines = []
    for key in EXPECTED_TOKEN_KEYS:
        label = CANONICAL_LABELS[key]
        if TOKEN_KINDS[key] == "header":
            lines.append((key, label))
        else:
            lines.append((key, f"{label} {fields[key]}"))
    return lines


def set_rtl_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_rtl_run(run, body=True):
    r_pr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    r_pr.append(rtl)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:bidi"), "ar-SA")
    r_pr.append(lang)
    if body:
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for attr in ("ascii", "hAnsi", "cs"):
            r_fonts.set(qn(f"w:{attr}"), "Arial")
        size = OxmlElement("w:szCs")
        size.set(qn("w:val"), "28")
        r_pr.append(size)
        run.font.name = "Arial"
        run.font.size = Pt(14)


def add_formatted_paragraph(document, text, style="Normal"):
    paragraph = document.add_paragraph(style=style)
    set_rtl_paragraph(paragraph)
    if style != "Heading 2":
        paragraph.paragraph_format.line_spacing = Pt(28)
    parts = text.split("\n")
    for index, part in enumerate(parts):
        run = paragraph.add_run(part)
        set_rtl_run(run, body=style != "Heading 2")
        if index < len(parts) - 1:
            run.add_break()
    return paragraph


def write_scripts_docx(path, lines=None, topic_id=1501):
    document = Document()
    add_formatted_paragraph(document, f"Script {topic_id}", "Heading 2")
    for _key, text in lines or token_lines():
        add_formatted_paragraph(document, text)
    document.save(path)


def write_thumbnail_docx(path, value=None, topic_id=1501):
    document = Document()
    add_formatted_paragraph(document, f"Script {topic_id}", "Heading 2")
    add_formatted_paragraph(
        document,
        value if value is not None else valid_fields()["youtube_thumbnail_1"],
    )
    document.save(path)


class FakeContext:
    def __init__(self, input_dir, output_dir, topic_ids=None):
        self.input_dir = str(input_dir)
        self.output_dir = str(output_dir)
        self.topic_ids = topic_ids
        self.results = {}

    def input_path(self, filename):
        return os.path.join(self.input_dir, filename)

    def output_path(self, filename):
        return os.path.join(self.output_dir, filename)

    def resolve(self, value):
        if not isinstance(value, str):
            return value
        if value.startswith("{") and value.endswith("}"):
            return self.results[value[1:-1]]
        return value


class LongTextReviewerTests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.input_dir = self.root / "input"
        self.output_dir = self.root / "output"
        self.input_dir.mkdir()
        self.output_dir.mkdir()
        (self.input_dir / "topics.json").write_text(
            json.dumps(
                {
                    "total_count": 1,
                    "titles": [{"id": 1501, "title": TITLE}],
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        write_scripts_docx(self.input_dir / "scripts_output.docx")
        write_thumbnail_docx(self.input_dir / "thumbnail_texts.docx")
        self.ctx = FakeContext(self.input_dir, self.output_dir, [1501])

    def tearDown(self):
        self.temp.cleanup()

    def build_cards(self, strict=True):
        return action_long_text_review_build_cards(
            {
                "id": "cards",
                "strict": strict,
                "source_file": "scripts_output.docx",
                "thumbnail_file": "thumbnail_texts.docx",
                "topics_file": "topics.json",
            },
            self.ctx,
        )

    def test_complete_valid_output_passes_deterministic_review(self):
        cards = self.build_cards(strict=True)
        self.assertTrue(cards["structure_report"]["structure_passed"])
        self.assertEqual(0, cards["structure_report"]["error_count"])
        self.assertEqual(list(FIELD_KEYS), list(cards["topics"]["1501"]["fields"]))
        self.assertTrue((self.output_dir / "long_text_structure_report.json").exists())

    def test_every_missing_token_is_detected(self):
        for missing_key in EXPECTED_TOKEN_KEYS:
            with self.subTest(missing_key=missing_key):
                lines = [item for item in token_lines() if item[0] != missing_key]
                entry = {
                    "id": "1501",
                    "paragraphs": [
                        {"index": index, "text": text}
                        for index, (_key, text) in enumerate(lines, start=1)
                    ],
                }
                issues = []
                _parse_script_fields(
                    entry,
                    "scripts_output.docx",
                    issues,
                    set(),
                )
                self.assertIn("TOKEN_MISSING", {item["code"] for item in issues})
                self.assertIn("TOKEN_SEQUENCE_WRONG", {item["code"] for item in issues})

    def test_every_adjacent_swap_is_detected(self):
        original = token_lines()
        for index in range(len(original) - 1):
            with self.subTest(index=index):
                lines = list(original)
                lines[index], lines[index + 1] = lines[index + 1], lines[index]
                entry = {
                    "id": "1501",
                    "paragraphs": [
                        {"index": number, "text": text}
                        for number, (_key, text) in enumerate(lines, start=1)
                    ],
                }
                issues = []
                _parse_script_fields(
                    entry,
                    "scripts_output.docx",
                    issues,
                    set(),
                )
                self.assertIn("TOKEN_SEQUENCE_WRONG", {item["code"] for item in issues})

    def test_every_duplicate_token_is_detected(self):
        original = token_lines()
        for index, (key, _text) in enumerate(original):
            with self.subTest(key=key):
                lines = list(original)
                lines.insert(index + 1, original[index])
                entry = {
                    "id": "1501",
                    "paragraphs": [
                        {"index": number, "text": text}
                        for number, (_key, text) in enumerate(lines, start=1)
                    ],
                }
                issues = []
                _parse_script_fields(
                    entry,
                    "scripts_output.docx",
                    issues,
                    set(),
                )
                self.assertIn("TOKEN_DUPLICATE", {item["code"] for item in issues})
                self.assertIn("TOKEN_SEQUENCE_WRONG", {item["code"] for item in issues})

    def test_missing_topic_and_thumbnail_mismatch_are_reported(self):
        self.ctx.topic_ids = [1501, 1502]
        cards = self.build_cards(strict=False)
        codes = {item["code"] for item in cards["issues"]}
        self.assertIn("SCRIPT_ID_MISSING", codes)
        self.assertIn("SCRIPT_ID_ORDER_WRONG", codes)
        write_thumbnail_docx(
            self.input_dir / "thumbnail_texts.docx",
            "نص مختلف تماما\nعن النص المطلوب",
        )
        cards = self.build_cards(strict=False)
        self.assertIn(
            "THUMBNAIL_OUTPUT_MISMATCH",
            {item["code"] for item in cards["issues"]},
        )

    def test_format_and_changed_label_errors_are_reported(self):
        document = Document()
        document.add_paragraph("Script 1501")
        for key, text in token_lines():
            if key == "youtube_description_1":
                text = text.replace("وصف الفيديو الأول:", "الوصف الأول:")
            document.add_paragraph(text)
        document.save(self.input_dir / "scripts_output.docx")
        cards = self.build_cards(strict=False)
        codes = {item["code"] for item in cards["issues"]}
        self.assertIn("LABEL_CHANGED", codes)
        self.assertIn("DOCX_PARAGRAPH_STYLE_WRONG", codes)
        self.assertIn("DOCX_PARAGRAPH_RTL_MISSING", codes)
        self.assertIn("DOCX_RUN_RTL_MISSING", codes)

    def valid_verdict(self, cards):
        topic = cards["topics"]["1501"]

        def evidence(source):
            return " ".join(str(source).replace("\n", " ").split()[:3])

        policy_sha256 = cards.get("review_policy_sha256") or _policy_sha256(
            "راجع الدقة والمعنى والترجمة."
        )
        cards["review_policy_sha256"] = policy_sha256
        review_request_id = _topic_review_request_id(
            "1501",
            topic,
            policy_sha256,
        )

        return {
            "schema_version": LONG_TEXT_REVIEW_SCHEMA_VERSION,
            "policy_sha256": policy_sha256,
            "request_id": review_request_id,
            "topic_id": 1501,
            "field_reviews": [
                {
                    "field": key,
                    "status": "سليم",
                    "evidence": evidence(topic["fields"][key]),
                    "reason": "",
                    "error_codes": [],
                }
                for key in FIELD_KEYS
            ],
            "cross_checks": [
                {
                    "check": key,
                    "status": "سليم",
                    "evidence": evidence(topic["title"]),
                    "reason": "",
                    "error_codes": [],
                }
                for key in SEMANTIC_CROSS_CHECKS
            ],
        }

    def test_semantic_prompt_and_valid_verdict_are_auditable(self):
        cards = self.build_cards(strict=True)
        self.ctx.results["cards"] = cards
        prompts = action_long_text_review_build_prompts(
            {
                "id": "prompts",
                "input": "{cards}",
                "instructions": "راجع الدقة والمعنى والترجمة.",
            },
            self.ctx,
        )
        self.assertEqual(1, len(prompts))
        self.assertIn("<BEGIN_UNTRUSTED_LONG_TEXT_DATA>", prompts[0])
        verdict = self.valid_verdict(cards)
        self.ctx.results["verdicts"] = [json.dumps(verdict, ensure_ascii=False)]
        report = action_long_text_review_parse_verdicts(
            {
                "id": "report",
                "input": "{verdicts}",
                "cards": "{cards}",
                "fail_incomplete": True,
            },
            self.ctx,
        )
        self.assertTrue(report["review_completed"])
        self.assertTrue(report["all_clear"])
        self.assertEqual(0, report["response_error_count"])

    def test_semantic_response_rejects_missing_item_bad_hash_and_fake_evidence(self):
        cards = self.build_cards(strict=True)
        self.ctx.results["cards"] = cards
        verdict = self.valid_verdict(cards)
        verdict["request_id"] = "wrong"
        verdict["field_reviews"].pop()
        verdict["field_reviews"][0]["evidence"] = "اقتباس غير موجود بالمصدر"
        self.ctx.results["verdicts"] = [json.dumps(verdict, ensure_ascii=False)]
        report = action_long_text_review_parse_verdicts(
            {
                "id": "report",
                "input": "{verdicts}",
                "cards": "{cards}",
                "fail_incomplete": False,
            },
            self.ctx,
        )
        self.assertFalse(report["review_completed"])
        messages = " | ".join(item["message"] for item in report["response_errors"])
        self.assertIn("request_id", messages)
        self.assertIn("field_reviews", messages)
        self.assertIn("اقتباسا حرفيا", messages)

    def test_semantic_finding_marks_only_affected_topic_for_regeneration(self):
        cards = self.build_cards(strict=True)
        self.ctx.results["cards"] = cards
        verdict = self.valid_verdict(cards)
        verdict["field_reviews"][0].update(
            {
                "status": "خطأ",
                "reason": "العنوان يحتوي ادعاء غير دقيق",
                "error_codes": ["SCIENTIFIC_ERROR"],
            }
        )
        self.ctx.results["verdicts"] = [json.dumps(verdict, ensure_ascii=False)]
        report = action_long_text_review_parse_verdicts(
            {
                "id": "report",
                "input": "{verdicts}",
                "cards": "{cards}",
                "fail_incomplete": True,
            },
            self.ctx,
        )
        self.assertTrue(report["review_completed"])
        self.assertFalse(report["all_clear"])
        self.assertEqual([1501], report["regenerate_topics"])
        payload = json.loads(
            (self.output_dir / "long_text_topics_to_regenerate.json").read_text(
                encoding="utf-8"
            )
        )
        self.assertEqual(1501, payload["topics"][0]["topic_id"])


if __name__ == "__main__":
    unittest.main()
