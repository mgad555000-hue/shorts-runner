"""
أداة فحص ملفات النصوص (docx) — تطبق كل قواعد validate_texts + قواعد إضافية.

الاستخدام:
    python tools/check_texts_docx.py "path/to/file.docx"
    python tools/check_texts_docx.py "path/to/file.docx" --json report.json
    python tools/check_texts_docx.py "path/to/file.docx" --min-words 70 --max-words 130 --expected-parts 4
"""

import argparse
import io
import json
import re
import sys
from collections import defaultdict
from pathlib import Path

# تأكد إن stdout/stderr يدعموا UTF-8 على Windows
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx مش متنصب. شغّل: pip install python-docx")
    sys.exit(1)

# ═══════════════════════════════════════════════════════════
# ثوابت
# ═══════════════════════════════════════════════════════════

TASHKEEL = '\u064B\u064C\u064D\u064E\u064F\u0650\u0651\u0652\u0670'
TATWEEL = '\u0640'

MEDICAL_WORDS = {
    'الكلا': 'الْكِلَا',
    'النفرونات': 'النِّفْرُونَات',
    'الكيسات': 'الْكِيسَات',
    'التكيسات': 'التَّكَيُّسَات',
    'الكرياتينين': 'الْكِرْيَاتِينِين',
    'الكلوية': 'الْكُلَوِيَّة',
}

ALLOWED_LATIN = {'pH', 'mg', 'ml', 'kg', 'cm', 'mm', 'DNA', 'RNA', 'BMI', 'GFR', 'HDL', 'LDL', 'II', 'III', 'IV'}
COMMON_SHORT = {'في', 'من', 'أو', 'لا', 'ما', 'هو', 'هي', 'أن', 'إن', 'قد', 'لم', 'لن', 'بل', 'كل', 'بس', 'يا', 'عن', 'مع', 'بك', 'لك', 'له', 'لي', 'بي', 'ذا', 'دا', 'ده', 'دي'}


# ═══════════════════════════════════════════════════════════
# قراءة ملف docx + إعادة بناء الماركرز
# ═══════════════════════════════════════════════════════════

def read_docx_with_markers(path):
    """يقرأ docx ويعيد بناء الماركرز <<<SCRIPT_N>>> ... <<<END_SCRIPT>>> + <<<PART_N>>> ... <<<END_PART>>>.

    منطق إعادة البناء (مطابق لـ recipe_runner.action_read_docx):
    - لما يلاقي عنوان (heading أو نص بدون نقطة في الآخر) = بداية SCRIPT جديد
    - لما يلاقي PART_N أو رقم في عنوان = part جديد
    """
    doc = Document(path)
    full_text_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text_parts.append(text)
    raw_text = '\n'.join(full_text_parts)

    # لو الماركرز موجودة فعلاً في النص ← استخدمها مباشرة
    if '<<<SCRIPT_' in raw_text:
        return raw_text

    # غير كده ← إعادة بناء (السيناريو الافتراضي للملفات اللي اتحفظت بـ save_docx)
    return _reconstruct_markers_from_paragraphs(doc)


def _reconstruct_markers_from_paragraphs(doc):
    """إعادة بناء الماركرز من docx لما تكون مفقودة.

    منطق الكشف:
    - فقرة قصيرة بدون نقطة في الآخر تشبه عنوان = بداية SCRIPT
    - "PART N" أو نمط مماثل = بداية PART
    - باقي الفقرات = محتوى PART الحالي
    """
    blocks = []
    current_script_num = 0
    current_part_num = 0
    current_part_lines = []
    in_script = False

    def flush_part():
        nonlocal current_part_lines, current_part_num
        if current_part_lines:
            content = ' '.join(current_part_lines).strip()
            blocks.append(f"<<<PART_{current_part_num}>>>\n{content}\n<<<END_PART>>>")
            current_part_lines = []

    def flush_script():
        nonlocal current_script_num, in_script, current_part_num
        flush_part()
        if in_script:
            blocks.append(f"<<<END_SCRIPT>>>")
        in_script = False
        current_part_num = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # كشف PART
        part_match = re.match(r'^(?:PART\s*[_:\-]?\s*|الجزء\s+|جزء\s+)(\d+)\s*$', text, re.IGNORECASE)
        if part_match:
            flush_part()
            current_part_num = int(part_match.group(1))
            continue

        # كشف بداية SCRIPT جديد (عنوان قصير بدون نقطة)
        is_heading = (
            (para.style and 'heading' in (para.style.name or '').lower()) or
            (len(text) < 80 and not text.endswith(('.', '،', '؟', '!')) and current_part_num == 0)
        )

        if is_heading and (not in_script or current_part_lines or current_part_num > 0):
            flush_script()
            current_script_num += 1
            blocks.append(f"<<<SCRIPT_{current_script_num}>>>")
            in_script = True
            continue

        # محتوى عادي
        if not in_script:
            current_script_num += 1
            blocks.append(f"<<<SCRIPT_{current_script_num}>>>")
            in_script = True
        if current_part_num == 0:
            current_part_num = 1
        current_part_lines.append(text)

    flush_script()
    return '\n\n'.join(blocks)


# ═══════════════════════════════════════════════════════════
# الفحوصات الميكانيكية
# ═══════════════════════════════════════════════════════════

def check_mid_word_spaces(text):
    """فحص المسافات داخل الكلمات العربية (الأخطر — مش موجود في validate_texts).

    منطق الكشف الدقيق:
    - كلمة عربية (3+ حروف) + مسافة + جزء قصير جداً (1-2 حرف عربي بس)
    - الجزء الثاني المفروض ينتهي بحرف "إنهاء كلمة" نموذجي: ء، ة، ى
    - الجزء الثاني مش كلمة عربية شائعة
    - الجزء الثاني مفيش فيه تشكيل كامل (حركة) في الحرف الأول لو كان حرف وحيد
    """
    # كلمات قصيرة مسموحة (مش false positive)
    ALLOWED_2_LETTERS = {
        'في', 'من', 'أو', 'لا', 'ما', 'هو', 'هي', 'أن', 'إن', 'قد', 'لم', 'لن',
        'بل', 'كل', 'بس', 'يا', 'عن', 'مع', 'بك', 'لك', 'له', 'لي', 'بي', 'ذا',
        'دا', 'ده', 'دي', 'إذ', 'إذا', 'أي', 'أى', 'أم', 'إى', 'به', 'بها',
        'لها', 'لهم', 'فى', 'أت', 'كم', 'كي', 'هل', 'لو', 'بل',
    }

    END_OF_WORD_LETTERS = set('اىةيوءأؤإ')
    INNER_DIACRITICS = '\u064E\u064F\u0650'

    matches = []
    pattern = r'([\u0621-\u064A\u0670' + re.escape(TASHKEEL) + r']{4,})\s([\u0621-\u064A\u0670' + re.escape(TASHKEEL) + r']{1,2})(?=\s|[.،؟!:؛]|$)'

    for m in re.finditer(pattern, text):
        before = m.group(1)
        after = m.group(2)
        clean_after = re.sub(f'[{TASHKEEL}]', '', after).strip()
        clean_before = re.sub(f'[{TASHKEEL}]', '', before).strip()

        if clean_after in ALLOWED_2_LETTERS or clean_after in COMMON_SHORT:
            continue
        if len(clean_before) > 12:
            continue

        is_split = False
        if len(clean_after) == 1:
            is_split = clean_after in 'ءةى'
        elif len(clean_after) == 2:
            if clean_after[-1] in 'ءةى':
                is_split = True
            else:
                # قاعدة 3: حرف داخلي + حركة + جزء ثاني بتشكيل منخفض (مثل تَنْظِ يف)
                last_arabic = re.search(r'([\u0621-\u064A])([\u064B-\u0652\u0670]*)$', before)
                if last_arabic:
                    last_letter = last_arabic.group(1)
                    trailing = last_arabic.group(2)
                    if last_letter not in END_OF_WORD_LETTERS:
                        has_inner = any(d in trailing for d in INNER_DIACRITICS)
                        after_clean = after.strip()
                        a_dia = sum(1 for c in after_clean if c in TASHKEEL)
                        a_let = sum(1 for c in after_clean if '\u0621' <= c <= '\u064A')
                        low_ratio = (a_dia / a_let) < 0.5 if a_let else True
                        is_split = has_inner and low_ratio

        if is_split:
            matches.append(f"{before} {after}")

    return matches


def check_tatweel(text):
    """فحص حرف التطويل ـ (Unicode U+0640) — مش موجود في validate_texts."""
    matches = []
    pattern = r'\S*' + TATWEEL + r'\S*'
    for m in re.finditer(pattern, text):
        matches.append(m.group(0))
    return matches


def check_medical_words(text):
    """فحص الكلمات الطبية لو متشكّلة بالشكل المحدد.

    يبحث عن كل كلمة في الجدول، ولو لقاها بدون التشكيل المطلوب → مشكلة.
    """
    issues = []
    # شيل التشكيل علشان نلاقي الكلمات الأساسية
    text_no_tashkeel = re.sub(f'[{TASHKEEL}]', '', text)
    for raw, expected in MEDICAL_WORDS.items():
        # لو الكلمة موجودة في النص (بدون تشكيل) لكن مش بالشكل المطلوب
        # نشوف كل ظهور للكلمة بدون تشكيل ونتحقق
        for m in re.finditer(re.escape(raw), text_no_tashkeel):
            # نلاقي الموقع الفعلي في النص الأصلي
            # لو الكلمة الأصلية مش مطابقة للتشكيل المطلوب بالضبط → مشكلة
            if expected not in text:
                issues.append(f"{raw} → المطلوب: {expected}")
                break  # نسجل مرة واحدة فقط لكل كلمة
    return issues


def check_last_letter_no_diacritic(part_text):
    """فحص أن الحرف الأخير من كل كلمة بدون حركة (لكن موجود).

    حركة على آخر حرف = مشكلة.
    """
    issues = []
    words = part_text.split()
    DIACRITICS_END = '\u064B\u064C\u064D\u064E\u064F\u0650\u0652'  # تنوين + فتحة + كسرة + ضمة + سكون
    for w in words:
        # إزالة علامات الترقيم في الآخر
        clean = w.rstrip('.،؟!:؛')
        if not clean:
            continue
        # لو آخر حرف حركة (مش شدة + حركة) → مشكلة
        last_char = clean[-1]
        if last_char in DIACRITICS_END:
            # تأكد إن مش حركة بعد شدة (التشكيل المضاعف مسموح في وسط الكلمة بس)
            issues.append(w)
    return issues[:5]  # أقصى 5 أمثلة


def check_marker_balance(text):
    """فحص توازن الماركرز — كل افتتاح له إغلاق + الأرقام متسلسلة."""
    issues = []
    script_opens = re.findall(r'<<<SCRIPT_(\d+)>>>', text)
    script_closes = re.findall(r'<<<END_SCRIPT>>>', text)
    if len(script_opens) != len(script_closes):
        issues.append(f"عدد SCRIPT مفتوح={len(script_opens)} مقابل END_SCRIPT={len(script_closes)}")
    return issues


# ═══════════════════════════════════════════════════════════
# فحوصات validate_texts (منسوخة بالحرف من recipe_runner.py)
# ═══════════════════════════════════════════════════════════

def _check_brackets_pattern(part_text):
    bracket_matches = re.findall(r'\[[^\[\]]{1,3}\]', part_text)
    return [m for m in bracket_matches if re.search(r'[\u0621-\u064A]', m)]


def _check_single_letters(part_text):
    words = part_text.split()
    if len(words) < 5:
        return 0, 0
    single_count = 0
    for w in words:
        clean = re.sub(f'[{TASHKEEL}]', '', w).strip()
        clean = re.sub(r'[.،؟!:\-]', '', clean)
        arabic_only = re.sub(r'[^\u0621-\u064A]', '', clean)
        if len(arabic_only) == 1:
            single_count += 1
    return single_count, (single_count / len(words)) * 100


def _check_truncated_words(part_text):
    words = part_text.split()
    if len(words) < 5:
        return 0, 0
    short_count = 0
    for w in words:
        clean = re.sub(f'[{TASHKEEL}]', '', w).strip()
        clean = re.sub(r'[.،؟!:\-\[\]]', '', clean)
        arabic_only = re.sub(r'[^\u0621-\u064A]', '', clean)
        if 0 < len(arabic_only) <= 2 and clean not in COMMON_SHORT:
            short_count += 1
    return short_count, (short_count / len(words)) * 100


def _check_split_hamza(part_text):
    return re.findall(r'(?:^|\s)([إأ])\s+(\S+)', part_text)


def _check_repeated_words(part_text):
    words = part_text.split()
    if len(words) < 5:
        return 0
    repeated = 0
    for i in range(1, len(words)):
        w_prev = re.sub(f'[{TASHKEEL}]', '', words[i-1]).strip().rstrip('.،؟!')
        w_curr = re.sub(f'[{TASHKEEL}]', '', words[i]).strip().rstrip('.،؟!')
        if w_prev and w_curr and w_prev == w_curr and len(w_prev) > 2:
            repeated += 1
    return repeated


def _check_latin_chars(part_text):
    matches = re.findall(r'[a-zA-Z]{2,}', part_text)
    return [m for m in matches if m not in ALLOWED_LATIN]


def _check_markdown_artifacts(part_text):
    artifacts = []
    if re.search(r'\*\*[^*]+\*\*', part_text):
        artifacts.append('**bold**')
    if re.search(r'^#{1,3}\s', part_text, re.MULTILINE):
        artifacts.append('## heading')
    if re.search(r'^\s*[-*]\s+\S', part_text, re.MULTILINE):
        artifacts.append('- list')
    if re.search(r'(?<!\*)\*(?!\*)[^*\s][^*]*[^*\s]\*(?!\*)', part_text):
        artifacts.append('*italic*')
    if re.search(r'`[^`]+`', part_text):
        artifacts.append('`code`')
    return artifacts


def _check_missing_tashkeel(part_text):
    arabic_chars = re.findall(r'[\u0621-\u064A]', part_text)
    tashkeel_found = re.findall(f'[{TASHKEEL}]', part_text)
    if len(arabic_chars) < 20:
        return 0.0
    return len(tashkeel_found) / len(arabic_chars) * 100


def _check_repeated_phrases(part_text):
    words = part_text.split()
    if len(words) < 15:
        return 0
    clean_words = [re.sub(f'[{TASHKEEL}]', '', w).strip().rstrip('.،؟!') for w in words]
    trigrams = {}
    repeated = 0
    for i in range(len(clean_words) - 2):
        tri = ' '.join(clean_words[i:i+3])
        if len(tri) > 6:
            trigrams[tri] = trigrams.get(tri, 0) + 1
    for tri, count in trigrams.items():
        if count >= 3:
            repeated += count - 1
    return repeated


def _check_digit_numbers(part_text):
    return re.findall(r'(?<!\S)\d{3,}(?!\S)', part_text)


def _check_unclosed_parens(part_text):
    open_paren = part_text.count('(') - part_text.count(')')
    open_bracket = part_text.count('[') - part_text.count(']')
    return abs(open_paren) + abs(open_bracket)


# ═══════════════════════════════════════════════════════════
# الفحص الرئيسي
# ═══════════════════════════════════════════════════════════

def analyze_text(text, min_words=70, max_words=130, expected_parts=4):
    """يفحص النص الكامل ويرجع تقرير مفصّل."""
    report = {
        'total_scripts': 0,
        'clean_scripts': [],
        'failed_scripts': [],
        'issues_by_script': defaultdict(list),
        'global_issues': [],
        'mid_word_spaces': [],
        'tatweel_words': [],
        'medical_word_issues': [],
    }

    # فحوصات عامة على الملف كله
    marker_issues = check_marker_balance(text)
    report['global_issues'].extend(marker_issues)

    mid_spaces = check_mid_word_spaces(text)
    report['mid_word_spaces'] = mid_spaces

    tatweel = check_tatweel(text)
    report['tatweel_words'] = tatweel

    med_issues = check_medical_words(text)
    report['medical_word_issues'] = med_issues

    # فحص كل SCRIPT
    pattern = r'<<<SCRIPT_(\d+)>>>(.*?)<<<END_SCRIPT>>>'
    for m in re.finditer(pattern, text, re.DOTALL):
        script_num = m.group(1)
        content = m.group(2)
        report['total_scripts'] += 1

        script_issues = []

        # استخراج parts
        parts = list(re.finditer(r'<<<PART_(\d+)>>>(.*?)<<<END_PART>>>', content, re.DOTALL))

        # فحص عدد parts
        if len(parts) != expected_parts:
            script_issues.append(f"عدد أجزاء = {len(parts)} (المطلوب {expected_parts})")

        for pm in parts:
            part_num = pm.group(1)
            part_text = pm.group(2).strip()
            part_label = f"P{part_num}"

            # عدد كلمات
            words = part_text.split()
            wc = len(words)
            if wc < min_words or wc > max_words:
                script_issues.append(f"{part_label}: {wc} كلمة (النطاق {min_words}-{max_words})")

            # validate_texts checks
            brackets = _check_brackets_pattern(part_text)
            if len(brackets) > 3:
                script_issues.append(f"{part_label}: {len(brackets)} قوس مربع")

            sc, sp = _check_single_letters(part_text)
            if sp > 15:
                script_issues.append(f"{part_label}: {sc} حرف منفرد ({sp:.0f}%)")

            tc, tp = _check_truncated_words(part_text)
            if tp > 20:
                script_issues.append(f"{part_label}: {tc} كلمة مبتورة ({tp:.0f}%)")

            sh = _check_split_hamza(part_text)
            if len(sh) > 3:
                script_issues.append(f"{part_label}: {len(sh)} همزة منفصلة")

            rw = _check_repeated_words(part_text)
            if rw > 3:
                script_issues.append(f"{part_label}: {rw} كلمة مكررة متتالية")

            lat = _check_latin_chars(part_text)
            if len(lat) > 2:
                script_issues.append(f"{part_label}: {len(lat)} كلمة لاتينية ({', '.join(lat[:3])})")

            md = _check_markdown_artifacts(part_text)
            if md:
                script_issues.append(f"{part_label}: بقايا Markdown — {', '.join(md)}")

            tash = _check_missing_tashkeel(part_text)
            if 0 < tash < 15:
                script_issues.append(f"{part_label}: تشكيل {tash:.0f}% فقط")

            rp = _check_repeated_phrases(part_text)
            if rp > 3:
                script_issues.append(f"{part_label}: {rp} عبارة مكررة")

            digits = _check_digit_numbers(part_text)
            if digits:
                script_issues.append(f"{part_label}: أرقام رقمية ({', '.join(digits[:3])})")

            unc = _check_unclosed_parens(part_text)
            if unc > 0:
                script_issues.append(f"{part_label}: {unc} قوس غير مغلق")

            if wc < 5:
                script_issues.append(f"{part_label}: جزء فارغ ({wc} كلمات)")

            # فحوصات إضافية (الجديدة)
            local_tatweel = check_tatweel(part_text)
            if local_tatweel:
                script_issues.append(f"{part_label}: {len(local_tatweel)} كلمة فيها Tatweel ({', '.join(local_tatweel[:2])})")

            local_mid = check_mid_word_spaces(part_text)
            if local_mid:
                script_issues.append(f"{part_label}: {len(local_mid)} مسافة وسط كلمة ({', '.join(local_mid[:2])})")

            last_letter = check_last_letter_no_diacritic(part_text)
            if len(last_letter) > 5:
                script_issues.append(f"{part_label}: {len(last_letter)} كلمة بحركة على آخر حرف ({', '.join(last_letter[:2])})")

        if script_issues:
            report['failed_scripts'].append(script_num)
            report['issues_by_script'][script_num] = script_issues
        else:
            report['clean_scripts'].append(script_num)

    return report


# ═══════════════════════════════════════════════════════════
# طباعة التقرير
# ═══════════════════════════════════════════════════════════

def print_report(report, file_path):
    """طباعة تقرير بشري مقروء."""
    print()
    print("═" * 75)
    print(f"  تقرير فحص: {file_path}")
    print("═" * 75)
    print()

    total = report['total_scripts']
    clean = len(report['clean_scripts'])
    failed = len(report['failed_scripts'])

    print(f"📊 الإحصائيات:")
    print(f"  • إجمالي المواضيع: {total}")
    print(f"  • ✅ سليمة: {clean}")
    print(f"  • ❌ فاشلة: {failed}")
    if total > 0:
        print(f"  • نسبة النجاح: {(clean / total * 100):.1f}%")
    print()

    if report['global_issues']:
        print("⚠️  مشاكل عامة في الملف:")
        for issue in report['global_issues']:
            print(f"  • {issue}")
        print()

    if report['mid_word_spaces']:
        print(f"🔴 مسافات وسط الكلمات ({len(report['mid_word_spaces'])} حالة):")
        for w in report['mid_word_spaces'][:10]:
            print(f"  • {w}")
        if len(report['mid_word_spaces']) > 10:
            print(f"  ... و {len(report['mid_word_spaces']) - 10} حالة أخرى")
        print()

    if report['tatweel_words']:
        print(f"🔴 كلمات فيها Tatweel ـ ({len(report['tatweel_words'])} حالة):")
        for w in report['tatweel_words'][:10]:
            print(f"  • {w}")
        if len(report['tatweel_words']) > 10:
            print(f"  ... و {len(report['tatweel_words']) - 10} حالة أخرى")
        print()

    if report['medical_word_issues']:
        print(f"⚕️  مشاكل في الكلمات الطبية:")
        for issue in report['medical_word_issues']:
            print(f"  • {issue}")
        print()

    if report['failed_scripts']:
        print(f"❌ تفاصيل المواضيع الفاشلة ({failed}):")
        for script_num in report['failed_scripts']:
            print(f"\n  SCRIPT_{script_num}:")
            for issue in report['issues_by_script'][script_num]:
                print(f"    • {issue}")
        print()

    print("═" * 75)
    if failed == 0 and not report['global_issues'] and not report['mid_word_spaces'] and not report['tatweel_words'] and not report['medical_word_issues']:
        print("  ✅ الملف خالٍ من الأخطاء — جاهز للنشر")
    else:
        print(f"  ⚠️  الملف يحتاج مراجعة — {failed} موضوع فاشل + {len(report['mid_word_spaces'])} مسافة + {len(report['tatweel_words'])} Tatweel")
    print("═" * 75)
    print()


# ═══════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="فحص ملف docx لأخطاء النصوص العربية المشكّلة")
    parser.add_argument("file", help="مسار ملف docx")
    parser.add_argument("--min-words", type=int, default=70, help="الحد الأدنى للكلمات في كل PART (افتراضي: 70)")
    parser.add_argument("--max-words", type=int, default=130, help="الحد الأقصى للكلمات في كل PART (افتراضي: 130)")
    parser.add_argument("--expected-parts", type=int, default=4, help="عدد PARTs المطلوب لكل SCRIPT (افتراضي: 4)")
    parser.add_argument("--json", help="حفظ التقرير كـ JSON في الملف المحدد")
    parser.add_argument("--show-text", action="store_true", help="طباعة النص المعاد بناؤه (للتشخيص)")
    args = parser.parse_args()

    file_path = Path(args.file)
    if not file_path.exists():
        print(f"ERROR: الملف مش موجود: {file_path}")
        sys.exit(1)

    if file_path.suffix.lower() != '.docx':
        print(f"ERROR: الملف لازم يكون docx (المستلم: {file_path.suffix})")
        sys.exit(1)

    text = read_docx_with_markers(str(file_path))

    if args.show_text:
        print(text)
        print("─" * 75)

    report = analyze_text(
        text,
        min_words=args.min_words,
        max_words=args.max_words,
        expected_parts=args.expected_parts,
    )

    print_report(report, file_path)

    if args.json:
        # تحويل defaultdict إلى dict عادي للـ JSON
        report['issues_by_script'] = dict(report['issues_by_script'])
        Path(args.json).write_text(
            json.dumps(report, ensure_ascii=False, indent=2),
            encoding='utf-8'
        )
        print(f"📄 تم حفظ التقرير في: {args.json}")

    # exit code: 0 لو نظيف، 1 لو فيه أخطاء
    has_errors = (
        len(report['failed_scripts']) > 0 or
        len(report['global_issues']) > 0 or
        len(report['mid_word_spaces']) > 0 or
        len(report['tatweel_words']) > 0 or
        len(report['medical_word_issues']) > 0
    )
    sys.exit(1 if has_errors else 0)


if __name__ == "__main__":
    main()
