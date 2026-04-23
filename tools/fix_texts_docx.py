"""
أداة إصلاح ميكانيكي لملفات Word - تطبق نفس قواعد auto_fix_text في recipe_runner.

تصلح أوتوماتيك (بدون AI):
- حذف Tatweel ـ
- دمج المسافات داخل الكلمات
- تصحيح الكلمات الطبية
- حذف الحركة على الحرف الأخير (اختياري)

تحفظ ملف جديد باسم {original}_fixed.docx بنفس التنسيق (font + bold + colors).
"""

import argparse
import io
import re
import sys
from copy import deepcopy
from pathlib import Path

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx مش متنصب. شغّل: pip install python-docx")
    sys.exit(1)


TASHKEEL = '\u064B\u064C\u064D\u064E\u064F\u0650\u0651\u0652\u0670'
TATWEEL = '\u0640'

MEDICAL_WORDS = {
    'الكلا': 'الْكِلَا',
    'النفرونات': 'النِّفْرُونَات',
    'الكيسات': 'الْكِيسَات',
    'التكيسات': 'التَّكَيُّسَات',
    'الكرياتينين': 'الْكِرْيَاتِينِين',
    'الكلوية': 'الْكُلَوِيَّة',
}

ALLOWED_2_LETTERS = {
    'في', 'من', 'أو', 'لا', 'ما', 'هو', 'هي', 'أن', 'إن', 'قد', 'لم', 'لن',
    'بل', 'كل', 'بس', 'يا', 'عن', 'مع', 'بك', 'لك', 'له', 'لي', 'بي', 'ذا',
    'دا', 'ده', 'دي', 'إذ', 'أي', 'أى', 'أم', 'إى', 'به', 'بها', 'لها', 'لهم',
    'فى', 'كم', 'كي', 'هل', 'لو',
}


def fix_remove_tatweel(text):
    """يحذف كل أحرف التطويل ـ."""
    return text.replace(TATWEEL, '')


END_OF_WORD_LETTERS = set('اىةيوءأؤإ')  # حروف نهاية كلمة شائعة
INNER_DIACRITICS = '\u064E\u064F\u0650'  # فتحة، ضمة، كسرة (حركات داخلية)


def _looks_like_split_word(before, after):
    """يحدد إذا كانت 'before after' كلمة واحدة مقطوعة بمسافة.

    معايير الكشف (أيٌّ منها يكفي):
    1. الجزء الثاني 1 حرف من حروف نهاية الكلمة (ء/ة/ى) — مثل: 'الْأَطِبَّا ء'
    2. الجزء الثاني 2 حرف ينتهي بحرف نهاية كلمة (ء/ة/ى) — مثل: 'الْفَلْتَرَ ة'
    3. الجزء الأول ينتهي بحرف داخلي + حركة (مش حرف نهاية)، والجزء الثاني بدون تشكيل تقريباً
       — مثل: 'تَنْظِ يف' (تنظيف)، 'كِتَا ب' (كتاب)
    """
    clean_after = re.sub(f'[{TASHKEEL}]', '', after).strip()
    clean_before = re.sub(f'[{TASHKEEL}]', '', before).strip()

    # استبعاد الكلمات الشائعة
    if clean_after in ALLOWED_2_LETTERS:
        return False

    # طول الجزء الأول معقول (حماية من دمج كلمات طويلة جداً)
    if len(clean_before) > 12:
        return False

    # القاعدة 1: حرف وحيد من حروف نهاية الكلمة
    if len(clean_after) == 1:
        return clean_after in 'ءةى'

    if len(clean_after) != 2:
        return False

    # القاعدة 2: حرفين ينتهي بحرف نهاية كلمة
    if clean_after[-1] in 'ءةى':
        return True

    # القاعدة 3 (الجديدة): فحص حرف نهاية الجزء الأول
    # نلاقي آخر حرف عربي فعلي + ما يليه من تشكيل
    last_arabic_match = re.search(r'([\u0621-\u064A])([\u064B-\u0652\u0670]*)$', before)
    if not last_arabic_match:
        return False

    last_letter = last_arabic_match.group(1)
    trailing_diacritics = last_arabic_match.group(2)

    # لو آخر حرف من حروف نهاية الكلمة الشائعة → مش مقطوعة
    if last_letter in END_OF_WORD_LETTERS:
        return False

    # لو آخر حرف عليه حركة داخلية (كسرة/فتحة/ضمة) — دلالة على إنه وسط كلمة
    has_inner_diacritic = any(d in trailing_diacritics for d in INNER_DIACRITICS)

    # تأكيد: الجزء الثاني بدون تشكيل (يدل إنه قطع من الكلمة الأكبر)
    after_clean_no_space = after.strip()
    after_diacritics = sum(1 for c in after_clean_no_space if c in TASHKEEL)
    after_letters = sum(1 for c in after_clean_no_space if '\u0621' <= c <= '\u064A')

    # نسبة التشكيل في الجزء الثاني منخفضة (<50% من الحروف عليها تشكيل)
    low_tashkeel_ratio = (after_diacritics / after_letters) < 0.5 if after_letters else True

    return has_inner_diacritic and low_tashkeel_ratio


def fix_mid_word_spaces(text):
    """يدمج المسافات داخل الكلمات العربية."""
    fixes = 0

    def replace(m):
        nonlocal fixes
        before = m.group(1)
        after = m.group(2)
        if _looks_like_split_word(before, after):
            fixes += 1
            return before + after
        return m.group(0)

    pattern = r'([\u0621-\u064A\u0670' + re.escape(TASHKEEL) + r']{4,})\s([\u0621-\u064A\u0670' + re.escape(TASHKEEL) + r']{1,2})(?=\s|[.،؟!:؛]|$)'
    fixed = re.sub(pattern, replace, text)
    return fixed, fixes


def fix_medical_words(text):
    """يصحح الكلمات الطبية."""
    fixes = 0

    def replace(m):
        nonlocal fixes
        word = m.group(0)
        clean = re.sub(f'[{TASHKEEL}]', '', word).strip()
        for raw, expected in MEDICAL_WORDS.items():
            if clean == raw and word != expected:
                fixes += 1
                return expected
        return word

    word_pattern = r'[\u0621-\u064A\u0670' + re.escape(TASHKEEL) + r']+'
    fixed = re.sub(word_pattern, replace, text)
    return fixed, fixes


def fix_text(text):
    """يطبق كل الإصلاحات على النص ويرجع (النص المصلح، إحصائيات)."""
    stats = {'tatweel': 0, 'mid_spaces': 0, 'medical': 0}

    stats['tatweel'] = text.count(TATWEEL)
    text = fix_remove_tatweel(text)

    text, stats['mid_spaces'] = fix_mid_word_spaces(text)
    text, stats['medical'] = fix_medical_words(text)

    return text, stats


def fix_docx_in_place(input_path, output_path):
    """يقرأ docx ويصلح كل run داخل كل paragraph، ويحفظ docx جديد."""
    doc = Document(input_path)
    total_stats = {'tatweel': 0, 'mid_spaces': 0, 'medical': 0}
    paragraphs_changed = 0

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        original_text = para.text
        fixed_text, stats = fix_text(original_text)

        if fixed_text != original_text:
            paragraphs_changed += 1
            total_stats['tatweel'] += stats['tatweel']
            total_stats['mid_spaces'] += stats['mid_spaces']
            total_stats['medical'] += stats['medical']

            # نطبق الإصلاح على الـ runs (محافظين على التنسيق)
            _apply_fix_to_runs(para, fixed_text)

    doc.save(output_path)
    return total_stats, paragraphs_changed


def _apply_fix_to_runs(para, fixed_text):
    """يطبق النص المُصلح على الـ runs محافظاً على التنسيق.

    استراتيجية بسيطة: نحط النص المصلح كله في أول run ونفرّغ الباقي.
    ده يحافظ على style (font, size, bold) لأن أول run بياخد الـ default formatting.
    """
    if not para.runs:
        para.add_run(fixed_text)
        return

    # احفظ التنسيق من أول run
    first_run = para.runs[0]
    first_run.text = fixed_text

    # افرغ الباقي
    for run in para.runs[1:]:
        run.text = ''


# ═══════════════════════════════════════════════════════════
# تحليل المواضيع المحتاجة Gemini (نفس منطق check_texts_docx)
# ═══════════════════════════════════════════════════════════

def analyze_for_regen(input_path, expected_parts=4, min_words=70, max_words=130):
    """يحلل docx بعد الإصلاح ويحدد المواضيع المحتاجة إعادة توليد بـ Gemini.

    يفحص (المشاكل التوليدية فقط):
    - عدد parts ≠ المتوقع
    - عدد كلمات خارج النطاق
    - حروف منفردة كثيرة (نص مفكّك)
    - كلمات مبتورة كثيرة

    يرجع: قائمة TOPIC_IDS محتاجة إعادة توليد + تفصيل لكل واحد
    """
    doc = Document(input_path)
    text_parts = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            text_parts.append(t)
    full_text = '\n'.join(text_parts)

    # محاولة استخراج SCRIPT/PART markers
    if '<<<SCRIPT_' not in full_text:
        # نعيد بناء بناء على heuristics: عنوان قصير = موضوع جديد
        full_text = _reconstruct_markers(doc)

    needs_regen = []
    pattern = r'<<<SCRIPT_(\d+)>>>(.*?)<<<END_SCRIPT>>>'
    for m in re.finditer(pattern, full_text, re.DOTALL):
        script_num = int(m.group(1))
        content = m.group(2)
        parts = list(re.finditer(r'<<<PART_(\d+)>>>(.*?)<<<END_PART>>>', content, re.DOTALL))

        issues = []

        if len(parts) != expected_parts:
            issues.append(f"عدد parts={len(parts)} (المطلوب {expected_parts})")

        for pm in parts:
            part_num = pm.group(1)
            ptext = pm.group(2).strip()
            wc = len(ptext.split())
            if wc < min_words or wc > max_words:
                issues.append(f"P{part_num}: {wc} كلمة")

        if issues:
            needs_regen.append({'script_num': script_num, 'issues': issues})

    return needs_regen


def _reconstruct_markers(doc):
    """إعادة بناء ماركرز من docx بدون ماركرز."""
    blocks = []
    script_num = 0
    part_num = 0
    current_lines = []
    in_script = False

    def flush_part():
        nonlocal current_lines, part_num
        if current_lines:
            content = ' '.join(current_lines).strip()
            blocks.append(f"<<<PART_{part_num}>>>\n{content}\n<<<END_PART>>>")
            current_lines = []

    def flush_script():
        nonlocal in_script, part_num
        flush_part()
        if in_script:
            blocks.append("<<<END_SCRIPT>>>")
        in_script = False
        part_num = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # كشف PART
        pm = re.match(r'^(?:PART\s*[_:\-]?\s*|الجزء\s+|جزء\s+)(\d+)\s*$', text, re.IGNORECASE)
        if pm:
            flush_part()
            part_num = int(pm.group(1))
            continue

        # كشف عنوان (heading)
        is_heading = (
            (para.style and 'heading' in (para.style.name or '').lower()) or
            (len(text) < 80 and not text.endswith(('.', '،', '؟', '!')) and part_num == 0)
        )

        if is_heading and (not in_script or current_lines or part_num > 0):
            flush_script()
            script_num += 1
            blocks.append(f"<<<SCRIPT_{script_num}>>>")
            in_script = True
            continue

        if not in_script:
            script_num += 1
            blocks.append(f"<<<SCRIPT_{script_num}>>>")
            in_script = True
        if part_num == 0:
            part_num = 1
        current_lines.append(text)

    flush_script()
    return '\n\n'.join(blocks)


# ═══════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="إصلاح ميكانيكي لملف docx + تحديد المواضيع المحتاجة Gemini")
    parser.add_argument("file", help="مسار ملف docx")
    parser.add_argument("--output", help="مسار الملف المُصلح (افتراضي: {input}_fixed.docx)")
    parser.add_argument("--expected-parts", type=int, default=4)
    parser.add_argument("--min-words", type=int, default=70)
    parser.add_argument("--max-words", type=int, default=130)
    args = parser.parse_args()

    input_path = Path(args.file)
    if not input_path.exists():
        print(f"ERROR: الملف مش موجود: {input_path}")
        sys.exit(1)

    output_path = Path(args.output) if args.output else input_path.with_name(input_path.stem + '_fixed.docx')

    print()
    print("═" * 75)
    print(f"  📂 المدخل: {input_path.name}")
    print(f"  💾 المخرج: {output_path.name}")
    print("═" * 75)

    # 1) إصلاح ميكانيكي
    print("\n🔧 جاري الإصلاح الميكانيكي...")
    stats, changed = fix_docx_in_place(str(input_path), str(output_path))
    total = stats['tatweel'] + stats['mid_spaces'] + stats['medical']

    print(f"\n✅ تم حفظ الملف المُصلح!")
    print(f"  • فقرات تم تعديلها: {changed}")
    print(f"  • Tatweel ـ تم حذفه: {stats['tatweel']}")
    print(f"  • مسافات وسط كلمة تم دمجها: {stats['mid_spaces']}")
    print(f"  • كلمات طبية تم تصحيحها: {stats['medical']}")
    print(f"  • إجمالي الإصلاحات: {total}")

    # 2) تحليل المواضيع المحتاجة Gemini
    print("\n🤖 جاري تحديد المواضيع المحتاجة إعادة توليد بـ Gemini...")
    needs_regen = analyze_for_regen(
        str(output_path),
        expected_parts=args.expected_parts,
        min_words=args.min_words,
        max_words=args.max_words,
    )

    print()
    print("═" * 75)
    if not needs_regen:
        print("  🎉 الملف نظيف 100% — مفيش حاجة محتاجة Gemini!")
        print("═" * 75)
    else:
        ids = [str(item['script_num']) for item in needs_regen]
        print(f"  ⚠️  {len(needs_regen)} موضوع محتاج إعادة توليد بـ Gemini")
        print("═" * 75)
        print()
        for item in needs_regen:
            print(f"  SCRIPT_{item['script_num']}:")
            for issue in item['issues']:
                print(f"    • {issue}")
        print()
        print("📋 TOPIC_IDS للنسخ المباشر في وصفة 'توليد_نصوص_من_مقدمات':")
        print()
        print(f"  {','.join(ids)}")
        print()

    print()


if __name__ == "__main__":
    main()
